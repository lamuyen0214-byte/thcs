import streamlit as st
from ai_config import get_ai_client # Lấy client từ file cấu hình chuẩn
from docx import Document
from io import BytesIO
from pypdf import PdfReader

def render_quiz_generator():
    st.subheader("🎯 Trình Tạo Đề Kiểm Tra Tự Động")
    
    # Khu vực tải dữ liệu
    uploaded_file = st.file_uploader("Tải tài liệu tham khảo (PDF, DOCX, TXT):", type=['pdf', 'docx', 'txt'])
    text_content = st.text_area("Nội dung bài giảng / Yêu cầu cụ thể:", height=150)
    
    # Khu vực tùy chỉnh cấu hình đề
    col1, col2 = st.columns(2)
    with col1:
        num_mcq = st.number_input("Số câu trắc nghiệm:", min_value=0, value=20)
        num_essay = st.number_input("Số câu tự luận:", min_value=0, value=5)
    with col2:
        use_source = st.checkbox("Tuyệt đối bám sát tài liệu", value=True)
        include_digital = st.checkbox("Tích hợp Năng lực số (AI, Tư duy tính toán)", value=True)
    
    # Nút bấm tối ưu UI
    if st.button("🚀 Tạo đề ngay", type="primary", use_container_width=True):
        
        # 1. XỬ LÝ DỮ LIỆU ĐẦU VÀO TẬP TRUNG (Theo chuẩn thầy cung cấp)
        combined = ""
        
        if uploaded_file:
            if uploaded_file.name.endswith(".pdf"):
                reader = PdfReader(uploaded_file)
                combined += "\n".join(page.extract_text() or "" for page in reader.pages)
            elif uploaded_file.name.endswith(".docx"):
                doc = Document(uploaded_file)
                combined += "\n".join(p.text for p in doc.paragraphs)
            elif uploaded_file.name.endswith(".txt"):
                combined += uploaded_file.getvalue().decode("utf-8", errors="ignore")
                
        if text_content:
            combined += "\n" + text_content
            
        # Kiểm tra dữ liệu rỗng
        if not combined.strip():
            st.warning("⚠️ Chưa có dữ liệu để tạo đề. Vui lòng nhập nội dung hoặc tải tài liệu.")
            return

        # 2. KIỂM TRA ĐIỀU KIỆN API
        client = get_ai_client()
        if client is None:
            st.error("⚠️ Gemini Client chưa được khởi tạo. Hãy kiểm tra API Key ở Sidebar!")
            return
        if not hasattr(client, 'models'):
            st.error("⚠️ Lỗi cấu hình SDK: Không tìm thấy model service.")
            return

        # 3. GỌI API VỚI PROMPT CHUẨN GDPT 2018
        with st.spinner("AI đang phân tích tài liệu và biên soạn đề..."):
            prompt = f"""
            Bạn là chuyên gia biên soạn đề kiểm tra theo Chương trình GDPT 2018 cấp THCS.

            Yêu cầu bắt buộc:
            {"- Chỉ sử dụng nội dung trong tài liệu, tuyệt đối không tự bổ sung kiến thức ngoài." if use_source else "- Căn cứ vào tài liệu và có thể mở rộng kiến thức phù hợp với lứa tuổi."}
            - Sinh {num_mcq} câu trắc nghiệm khách quan.
            - Sinh {num_essay} câu tự luận.
            - Trình bày rõ CÂU HỎI.
            - Cung cấp ĐÁP ÁN chi tiết và HƯỚNG DẪN CHẤM.
            - Kèm theo MA TRẬN ĐỀ THI ở phần đầu.
            {"- Đặc biệt chú trọng lồng ghép các yếu tố Năng lực số, ứng dụng AI hoặc Tư duy tính toán vào một số câu hỏi thực tế." if include_digital else ""}

            Tài liệu tham khảo:
            {combined}
            """
            
            try:
                response = client.models.generate_content(
                    model="gemini-1.5-flash", 
                    contents=prompt
                )
                
                # Xử lý Safety Filter an toàn
                result = getattr(response, "text", "")
                if not result:
                    st.warning("⚠️ AI không trả về nội dung. Có thể nội dung bị chặn bởi bộ lọc an toàn của Google.")
                    return
                
                st.session_state.current_quiz = result
                st.rerun()
                
            except Exception as e:
                st.error(f"❌ Lỗi hệ thống AI: {str(e)}")

    # 4. HIỂN THỊ KẾT QUẢ VÀ XUẤT WORD
    if "current_quiz" in st.session_state:
        st.markdown("---")
        st.markdown(st.session_state.current_quiz)
        
        # Nút xuất Word chuyên nghiệp
        doc = Document()
        doc.add_heading("ĐỀ KIỂM TRA", level=1)
        doc.add_paragraph(st.session_state.current_quiz)
        
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        st.download_button(
            label="⬇️ Tải file Word (.docx)",
            data=bio,
            file_name="De_kiem_tra_GDPT2018.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
