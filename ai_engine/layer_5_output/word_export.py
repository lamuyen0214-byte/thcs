# File: ai_engine/layer_5_output/word_export.py
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io

class WordExportEngine:
    @staticmethod
    def export_to_word(content_markdown, title="Kich_Ban_STEM"):
        """
        Hàm xử lý xuất nội dung ra file Word (.docx).
        Trong tương lai, Lớp 4 (Formatting Layer) sẽ tiền xử lý các thẻ $latex$ 
        thành mã OMML trước khi đưa vào hàm này để đảm bảo công thức KHTN hiển thị chuẩn.
        """
        doc = docx.Document()
        
        # Thiết lập style mặc định
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Thêm tiêu đề
        heading = doc.add_heading(title.replace("_", " "), level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Xử lý đoạn văn bản (Tạm thời xử lý text thô, sẽ nâng cấp bộ phân tích Markdown sau)
        for line in content_markdown.split('\n'):
            if line.strip():
                if line.startswith('###'):
                    doc.add_heading(line.replace('###', '').strip(), level=3)
                elif line.startswith('##'):
                    doc.add_heading(line.replace('##', '').strip(), level=2)
                elif line.startswith('#'):
                    doc.add_heading(line.replace('#', '').strip(), level=1)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:].strip(), style='List Bullet')
                else:
                    doc.add_paragraph(line.strip())

        # Lưu vào bộ nhớ đệm (Buffer) để Streamlit cho phép tải xuống
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
