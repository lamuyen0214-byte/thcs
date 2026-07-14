import streamlit as st
import pandas as pd
import io
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import google.generativeai as genai

# 1. HÀM ĐỌC FILE WORD TẢI LÊN
def read_docx(file):
    try:
        doc = docx.Document(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        return f"Lỗi đọc file: {e}"

# 2. HÀM XUẤT FILE WORD (CHUẨN HÀNH CHÍNH)
def create_word_document(tieu_de, ngay_hop, dia_diem, chu_tri, thu_ky, vang_mat, noi_dung, ket_luan):
    doc = docx.Document()
    
    # Quốc hiệu - Tiêu ngữ
    p_quoc_hieu = doc.add_paragraph()
    p_quoc_hieu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_qh = p_quoc_hieu.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    run_qh.bold = True
    run_qh.font.size = Pt(13)
    run_tn = p_quoc_hieu.add_run("Độc lập - Tự do - Hạnh phúc")
    run_tn.bold = True
    run_tn.font.size = Pt(14)
    
    doc.add_paragraph() # Dòng trống
    
    # Tên biên bản
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("BIÊN BẢN CUỘC HỌP\n")
    run_title.bold = True
    run_title.font.size = Pt(16)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(f"({tieu_de})")
    run_sub.italic = True
    run_sub.font.size = Pt(13)
    
    doc.add_paragraph()
    
    # Thành phần
    doc.add_paragraph(f"Thời gian: {ngay_hop}")
    doc.add_paragraph(f"Địa điểm: {dia_diem}")
    doc.add_paragraph(f"Chủ trì: {chu_tri}")
    doc.add_paragraph(f"Thư ký: {thu_ky}")
    doc.add_paragraph(f"Vắng mặt: {vang_mat}")
    
    # Nội dung
    doc.add_heading('I. NỘI DUNG TRIỂN KHAI:', level=2)
    doc.add_paragraph(noi_dung)
    
    # Kết luận
    doc.add_heading('II. KẾT LUẬN CỦA CHỦ TỌA:', level=2)
    doc.add_paragraph(ket_luan)
    
    doc.add_paragraph()
    
    # Chữ ký (Dùng bảng ẩn viền để chia 2 bên)
    table = doc.add_table(rows=1, cols=2)
    cell_1 = table.cell(0, 0)
    cell_1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_1.paragraphs[0].add_run("THƯ KÝ\n(Ký và ghi rõ họ tên)").bold = True
    
    cell_2 = table.cell(0, 1)
    cell_2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_2.paragraphs[0].add_run("CHỦ TRÌ\n(Ký và ghi rõ họ tên)").bold = True
    
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# 3. GIAO DIỆN CHÍNH CỦA BIÊN BẢN
def render_bien_ban(supabase):
    st.markdown("## 📝 Quản Lý Hồ Sơ Biên Bản Tổ Chuyên Môn")
    bb_tabs = st.tabs(["➕ Viết Biên Bản AI", "📋 Kho Biên Bản"])
    
    # --- TAB TẠO BIÊN BẢN (TÍCH HỢP AI) ---
    with bb_tabs[0]:
        st.markdown("### ✍️ Thông tin cuộc họp")
        
        # Menu xổ xuống (Yêu cầu 2)
        loai_sinh_hoat = [
            "Sinh hoạt chuyên môn theo NCBH", 
            "Sinh hoạt theo chuyên đề", 
            "Sinh hoạt hành chính định kỳ", 
            "Sinh hoạt bồi dưỡng thường xuyên"
        ]
        tieu_de = st.selectbox("Chọn Chủ đề cuộc họp:", loai_sinh_hoat)
        
        c1, c2, c3 = st.columns(3)
        ngay_hop = c1.date_input("Ngày họp")
        chu_tri = c2.text_input("Chủ trì", value="Lê Hồng Dưỡng")
        thu_ky = c3.text_input("Thư ký", placeholder="Nhập tên thư ký...")
        
        c4, c5 = st.columns(2)
        dia_diem = c4.text_input("Địa điểm", value="Văn phòng trường")
        vang_mat = c5.text_input("Vắng mặt", placeholder="VD: Không")
        
        st.markdown("---")
        st.markdown("### 🤖 Trợ lý AI Hỗ trợ viết Biên bản")
        
        # Tải file Kế hoạch & Checkbox (Yêu cầu 3 & 4)
        file_ke_hoach = st.file_uploader("Tải lên 'Kế hoạch tổ CM' (Định dạng .docx) để AI bám sát", type=['docx'])
        bam_sat_kh = st.checkbox("🎯 Yêu cầu AI BẮT BUỘC bóc tách và bám sát nội dung file Kế hoạch tải lên", value=True)
        ghi_chu_ai = st.text_input("Ghi chú thêm cho AI (Tùy chọn)", placeholder="VD: Nhấn mạnh vào việc thi giáo viên giỏi...")

        # Nơi chứa kết quả AI trả về (Dùng session_state để lưu tạm)
        if 'ai_noi_dung' not in st.session_state: st.session_state['ai_noi_dung'] = ""
        if 'ai_ket_luan' not in st.session_state: st.session_state['ai_ket_luan'] = ""

        if st.button("✨ Nhờ AI Viết Biên Bản", type="primary"):
            # Lấy API Key từ thanh bên (giả sử thầy lưu ở st.session_state['gemini_api_key'])
            # Nếu chưa có biến đó, yêu cầu nhập
            api_key = st.session_state.get('gemini_api_key', "")
            if not api_key:
                st.error("⚠️ Thầy cần dán Gemini API Key ở thanh menu bên trái trước khi dùng tính năng này!")
            else:
                try:
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel('gemini-1.5-flash')
                    
                    # Xây dựng Prompt theo từng loại Menu (Yêu cầu 5)
                    prompt_base = f"Đóng vai một thư ký cuộc họp chuyên nghiệp của trường học. Hãy viết chi tiết phần [NỘI DUNG] và [KẾT LUẬN] cho biên bản '{tieu_de}'. "
                    
                    if tieu_de == "Sinh hoạt chuyên môn theo NCBH":
                        prompt_base += "Nội dung cần tập trung vào việc phân tích kế hoạch bài dạy, dự giờ, quan sát hoạt động của học sinh, rút kinh nghiệm và định hướng điều chỉnh phương pháp dạy học. "
                    elif tieu_de == "Sinh hoạt theo chuyên đề":
                        prompt_base += "Nội dung cần tập trung vào báo cáo lý thuyết chuyên đề, thảo luận thực tiễn áp dụng tại trường, và thống nhất các bước thực hiện chuyên đề. "
                    elif tieu_de == "Sinh hoạt hành chính định kỳ":
                        prompt_base += "Nội dung cần đánh giá chi tiết công tác tháng qua (ưu khuyết điểm), tình hình nề nếp học sinh, và phân công nhiệm vụ cụ thể cho tháng tới. "
                    elif tieu_de == "Sinh hoạt bồi dưỡng thường xuyên":
                        prompt_base += "Nội dung tập trung vào việc triển khai các module bồi dưỡng, chia sẻ phương pháp sư phạm, và báo cáo kết quả tự học. "
                    
                    # Xử lý file đính kèm
                    if file_ke_hoach and bam_sat_kh:
                        docx_text = read_docx(file_ke_hoach)
                        prompt_base += f"\n\nBẮT BUỘC: Bạn phải đọc và bóc tách các nội dung từ tài liệu 'Kế hoạch tổ CM' sau đây để viết biên bản (nghĩa là cuộc họp đã bàn luận và triển khai chính xác những gì kế hoạch này đề ra):\n[KẾ HOẠCH BẮT ĐẦU]\n{docx_text}\n[KẾ HOẠCH KẾT THÚC]\n"
                    
                    if ghi_chu_ai:
                        prompt_base += f"\nLưu ý thêm từ Tổ trưởng: {ghi_chu_ai}"
                        
                    prompt_base += "\n\nHãy trả về kết quả theo cấu trúc sau (Không dùng markdown in đậm quá nhiều, viết văn bản liền mạch):\nPHẦN 1: NỘI DUNG\n(viết chi tiết ở đây)\nPHẦN 2: KẾT LUẬN\n(viết chi tiết ở đây)"

                    with st.spinner("🤖 AI đang đọc tài liệu và soạn biên bản. Thầy đợi chút nhé..."):
                        response = model.generate_content(prompt_base)
                        ai_text = response.text
                        
                        # Tách phần Nội dung và Kết luận
                        try:
                            phan_1 = ai_text.split("PHẦN 2: KẾT LUẬN")[0].replace("PHẦN 1: NỘI DUNG", "").strip()
                            phan_2 = ai_text.split("PHẦN 2: KẾT LUẬN")[1].strip()
                            st.session_state['ai_noi_dung'] = phan_1
                            st.session_state['ai_ket_luan'] = phan_2
                        except:
                            st.session_state['ai_noi_dung'] = ai_text # Nếu AI trả về ko đúng form, cho hết vào nội dung
                            
                except Exception as e:
                    st.error(f"Lỗi kết nối AI: {e}")

        # Cho phép người dùng chỉnh sửa lại văn bản AI sinh ra
        noi_dung = st.text_area("1. Nội dung cuộc họp (Có thể chỉnh sửa):", value=st.session_state['ai_noi_dung'], height=250)
        ket_luan = st.text_area("2. Kết luận (Có thể chỉnh sửa):", value=st.session_state['ai_ket_luan'], height=150)
        
        if st.button("💾 LƯU BIÊN BẢN LÊN HỆ THỐNG", use_container_width=True):
            new_bb = {
                "tieu_de": tieu_de, "ngay_hop": str(ngay_hop), "dia_diem": dia_diem,
                "chu_tri": chu_tri, "thu_ky": thu_ky, "vang_mat": vang_mat if vang_mat else "Không",
                "noi_dung": noi_dung, "ket_luan": ket_luan
            }
            try:
                supabase.table("bien_ban").insert(new_bb).execute()
                st.success("🎉 Đã lưu biên bản thành công! Thầy sang tab 'Kho Biên Bản' để xem và tải file Word.")
                # Xóa session
                st.session_state['ai_noi_dung'] = ""
                st.session_state['ai_ket_luan'] = ""
            except Exception as e:
                st.error(f"Lỗi lưu Supabase: {e}")

    # --- TAB KHO BIÊN BẢN (XEM & TẢI FILE WORD) ---
    with bb_tabs[1]:
        try:
            response = supabase.table("bien_ban").select("*").order("ngay_hop", desc=True).execute()
            df_bb = pd.DataFrame(response.data)
        except:
            df_bb = pd.DataFrame()

        if df_bb.empty:
            st.info("📭 Tổ chuyên môn chưa có biên bản nào được lưu.")
        else:
            for index, row in df_bb.iterrows():
                with st.expander(f"📌 {row['ngay_hop']} | {row['tieu_de']}"):
                    st.markdown(f"**Chủ trì:** {row['chu_tri']} | **Thư ký:** {row['thu_ky']}")
                    st.markdown(f"**Nội dung:**\n{row['noi_dung']}")
                    st.markdown(f"**Kết luận:**\n{row['ket_luan']}")
                    
                    # Nút Tải file Word (Yêu cầu 6 & 7)
                    docx_file = create_word_document(
                        row['tieu_de'], row['ngay_hop'], row['dia_diem'], 
                        row['chu_tri'], row['thu_ky'], row['vang_mat'], 
                        row['noi_dung'], row['ket_luan']
                    )
                    
                    colA, colB = st.columns(2)
                    with colA:
                        st.download_button(
                            label="📥 Tải Biên Bản (File Word chuẩn)",
                            data=docx_file,
                            file_name=f"BienBan_{row['ngay_hop']}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_{row['id']}"
                        )
                    with colB:
                        if st.button("🗑️ Xóa biên bản", key=f"del_{row['id']}"):
                            supabase.table("bien_ban").delete().eq("id", row['id']).execute()
                            st.rerun()
