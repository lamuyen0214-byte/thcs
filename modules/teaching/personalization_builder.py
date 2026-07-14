import streamlit as st
import io
import docx
from ai_engine.ai_config import get_api_key
from ai_engine.ai_runner import run_ai_with_fallback

def render_personalization_module(api_key=""):
    # Tinh chỉnh CSS
    st.markdown("""
    <style>
    div[data-testid="stAppViewBlockContainer"], .main .block-container { padding-top: 3.5rem !important; padding-bottom: 3rem !important; }
    .stSelectbox label p, .stTextInput label p, .stTextArea label p { color: #0000FF !important; font-weight: 600 !important; font-size: 14px !important; }
    .stMarkdown p strong { color: #FF0000 !important; font-size: 15px !important; }
    .stButton>button { font-weight: bold; border-radius: 6px; }
    </style>
    """, unsafe_allow_html=True)

    # 1. GIAO DIỆN NHẬP HỒ SƠ HỌC SINH
    col1, col2, col3 = st.columns([1.5, 1, 1.5])
    with col1:
        ten_hs = st.text_input("👤 Tên học sinh (hoặc Tên nhóm)", placeholder="Ví dụ: Nguyễn Văn A...")
    with col2:
        mon_hoc = st.selectbox("📚 Môn học", ["KHTN", "Vật Lý", "Hóa Học", "Sinh Học", "Toán", "Ngữ văn", "Tiếng Anh"])
    with col3:
        muc_tieu = st.text_input("🎯 Mục tiêu hướng tới", placeholder="Ví dụ: Đạt 8 điểm cuối kỳ, Đậu chuyên Lý...")

    col_style, col_time = st.columns([1, 1])
    with col_style:
        phong_cach = st.selectbox("🧠 Phong cách tiếp thu tốt nhất", [
            "Hình ảnh trực quan (Visual - Thích xem sơ đồ, video)",
            "Thính giác (Auditory - Thích nghe giảng, thảo luận)",
            "Vận động (Kinesthetic - Thích thực hành, làm thí nghiệm)",
            "Đọc/Viết (Reading/Writing - Thích đọc sách, ghi chép)"
        ])
    with col_time:
        thoi_gian = st.selectbox("⏱️ Thời gian triển khai lộ trình", ["2 tuần", "1 tháng", "1 học kỳ", "Cả năm học"])

    col_manh, col_yeu = st.columns([1, 1])
    with col_manh:
        diem_manh = st.text_area("💪 Điểm mạnh / Năng lực hiện tại", placeholder="Ví dụ: Tính toán nhanh, tư duy logic tốt, chăm chỉ...", height=80)
    with col_yeu:
        diem_yeu = st.text_area("⚠️ Điểm yếu / Vị trí hổng kiến thức", placeholder="Ví dụ: Lười học thuộc lý thuyết, hay sai dấu, cẩu thả...", height=80)

    # 2. KHỞI TẠO AI SÁNG TẠO LỘ TRÌNH
    if st.button("🚀 AI THIẾT LẬP LỘ TRÌNH CÁ NHÂN HÓA", type="primary", use_container_width=True):
        if not ten_hs.strip() or not diem_yeu.strip():
            st.warning("Vui lòng nhập tên học sinh và ít nhất một điểm yếu/vấn đề cần khắc phục!")
            st.stop()
            
        final_key = api_key if api_key else get_api_key()
        
        prompt = f"""
Bạn là một Chuyên gia Giáo dục Cá nhân hóa (Personalized Learning) xuất sắc.
Hãy lập một "Kế hoạch Học tập Cá nhân hóa" (IEP) cho học sinh cấp THCS/THPT.
- Tên học sinh: {ten_hs}
- Môn học: {mon_hoc}
- Mục tiêu: {muc_tieu}
- Thời gian: {thoi_gian}
- Phong cách tiếp thu: {phong_cach}
- Điểm mạnh hiện có: {diem_manh}
- Điểm yếu/Hổng kiến thức: {diem_yeu}

Yêu cầu cấu trúc bản Kế hoạch (Trình bày bằng Markdown rõ ràng, chuyên nghiệp):
1. PHÂN TÍCH CHẨN ĐOÁN: Tóm tắt ngắn gọn tình trạng của học sinh dựa trên thông tin đã cung cấp.
2. CHIẾN LƯỢC SƯ PHẠM: Đề xuất phương pháp giảng dạy riêng biệt phù hợp nhất với 'Phong cách tiếp thu' và 'Điểm mạnh' của em này.
3. LỘ TRÌNH THỰC HIỆN ({thoi_gian}): Chia nhỏ kế hoạch theo từng giai đoạn (tuần/tháng). Mỗi giai đoạn cần có nhiệm vụ học tập cụ thể khắc phục 'Điểm yếu'.
4. PHƯƠNG PHÁP ĐÁNH GIÁ & ĐO LƯỜNG: Cách giáo viên kiểm tra xem học sinh có tiến bộ hay không (Bài test nhỏ, vấn đáp, sản phẩm thực hành...).
5. LỜI KHUYÊN TÂM LÝ: Đề xuất cách giáo viên động viên, khích lệ học sinh này.
"""
        with st.spinner("AI đang phân tích hồ sơ và lên chiến lược cá nhân hóa..."):
            try:
                result = run_ai_with_fallback(prompt=prompt, api_key=final_key, model_mode="flash")
                if result.get("success"):
                    st.session_state['current_personalized_plan'] = {"title": ten_hs, "content": result.get("text")}
                else:
                    st.error(f"❌ Lỗi AI: {result.get('error')}")
            except Exception as e:
                st.error(f"❌ Lỗi hệ thống: {str(e)}")

    # 3. KẾT QUẢ VÀ XUẤT FILE WORD
    if 'current_personalized_plan' in st.session_state:
        st.markdown("---")
        with st.expander("📄 XEM TRƯỚC LỘ TRÌNH CÁ NHÂN HÓA", expanded=True):
            st.markdown(st.session_state['current_personalized_plan']['content'])
            
            def export_plan_to_word(data):
                doc = docx.Document()
                doc.add_heading(f"KẾ HOẠCH HỌC TẬP CÁ NHÂN HÓA", 0)
                doc.add_heading(f"Học sinh: {data['title']}", 1)
                doc.add_paragraph(data['content'])
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer

            if st.button("📥 Tải Kế hoạch (Word)", use_container_width=True):
                word_file = export_plan_to_word(st.session_state['current_personalized_plan'])
                st.download_button("Xác nhận Tải về", data=word_file, file_name=f"Ke_Hoach_Ca_Nhan_{ten_hs.replace(' ', '_')}.docx", use_container_width=True)
