import streamlit as st
import io
import docx
from gtts import gTTS
from ai_engine.ai_config import get_api_key
from ai_engine.ai_runner import run_ai_with_fallback

def render_video_module(api_key=""):
    # Tinh chỉnh CSS
    st.markdown("""
    <style>
    div[data-testid="stAppViewBlockContainer"], .main .block-container { padding-top: 3.5rem !important; padding-bottom: 3rem !important; }
    .stSelectbox label p, .stTextArea label p { color: #0000FF !important; font-weight: 600 !important; font-size: 14px !important; }
    .stMarkdown p strong { color: #FF0000 !important; font-size: 15px !important; }
    .stButton>button { font-weight: bold; border-radius: 6px; }
    audio { width: 100%; margin-top: 10px; }
    </style>
    """, unsafe_allow_html=True)

    # 1. GIAO DIỆN NHẬP LIỆU KỊCH BẢN VIDEO
    col1, col2 = st.columns([2, 1])
    with col1:
        chu_de_video = st.text_input("🎬 Chủ đề Video cần tạo", placeholder="Ví dụ: Vòng tuần hoàn của nước, Chiến thắng Điện Biên Phủ, Cấu tạo nguyên tử...")
    with col2:
        thoi_luong = st.selectbox("⏳ Thời lượng mong muốn", ["Dưới 1 phút (Shorts/TikTok)", "Khoảng 2-3 phút", "Khoảng 5 phút"])

    phong_cach = st.selectbox("🎨 Phong cách kể chuyện", [
        "Trang trọng, chuẩn mực (Phim tài liệu khoa học)", 
        "Hài hước, gần gũi (Phù hợp Gen Z/Học sinh cấp 2)", 
        "Giọng điệu bí ẩn, kể chuyện (Kích thích tò mò)"
    ])
    
    noi_dung_chinh = st.text_area("📄 Các ý chính bắt buộc phải có trong Video", placeholder="Nhập các keyword hoặc dàn ý ngắn để AI bám sát...")

    # 2. AI SINH KỊCH BẢN PHÂN CẢNH (SCENE-BY-SCENE SCRIPT)
    if st.button("🚀 AI SÁNG TẠO KỊCH BẢN PHÂN CẢNH", type="primary", use_container_width=True):
        if not chu_de_video.strip():
            st.warning("Vui lòng nhập Chủ đề Video!")
            st.stop()
            
        final_key = api_key if api_key else get_api_key()
        
        prompt = f"""
Bạn là một Đạo diễn và Chuyên gia Viết Kịch bản Video Giáo dục xuất sắc.
Hãy viết một kịch bản Video bài giảng về chủ đề: "{chu_de_video}".
Thời lượng dự kiến: {thoi_luong}.
Phong cách kể chuyện: {phong_cach}.
Nội dung cốt lõi cần có: {noi_dung_chinh if noi_dung_chinh else "Tự do sáng tạo bám sát chủ đề"}

YÊU CẦU CẤU TRÚC KỊCH BẢN TRÌNH BÀY DƯỚI DẠNG BẢNG MARKDOWN:
| Cảnh (Scene) | Hình ảnh mô tả (Visual/Prompt tạo ảnh) | Lời bình đọc Voice (Audio/Voiceover) | Thời gian dự kiến |
| :--- | :--- | :--- | :--- |
| Cảnh 1 | ... | ... | ... |

Hãy chia nhỏ thành các cảnh ngắn (mỗi cảnh 5-10 giây). Lời bình (Voiceover) phải cực kỳ cuốn hút, tự nhiên đúng như phong cách đã chọn. Cột "Hình ảnh mô tả" hãy viết giống như một câu lệnh (Prompt) để giáo viên có thể dùng đưa vào các AI vẽ ảnh (như Midjourney/Bing Image Creator).
"""
        with st.spinner("Đạo diễn AI đang lên kịch bản từng giây cho video..."):
            try:
                result = run_ai_with_fallback(prompt=prompt, api_key=final_key, model_mode="flash")
                if result.get("success"):
                    st.session_state['current_video_script'] = {
                        "title": chu_de_video, 
                        "content": result.get("text"),
                        "voice_text": "" # Chuẩn bị biến lưu text để đọc giọng nói
                    }
                else:
                    st.error(f"❌ Lỗi AI: {result.get('error')}")
            except Exception as e:
                st.error(f"❌ Lỗi hệ thống: {str(e)}")

    # 3. KẾT QUẢ VÀ TÍNH NĂNG ĐỌC GIỌNG NÓI (TEXT-TO-SPEECH)
    if 'current_video_script' in st.session_state:
        st.markdown("---")
        with st.expander("📝 KỊCH BẢN PHÂN CẢNH VIDEO (SCENE SCRIPT)", expanded=True):
            st.markdown(st.session_state['current_video_script']['content'])
            
            # Xuất Kịch bản ra Word
            def export_script_to_word(data):
                doc = docx.Document()
                doc.add_heading(f"KỊCH BẢN VIDEO: {data['title'].upper()}", 0)
                doc.add_paragraph(data['content'])
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer

            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.button("📥 Tải Kịch bản (Word)", use_container_width=True):
                    word_file = export_script_to_word(st.session_state['current_video_script'])
                    st.download_button("Xác nhận Tải Kịch bản", data=word_file, file_name=f"Kich_Ban_{chu_de_video.replace(' ', '_')}.docx", use_container_width=True)
            
            # Công cụ sinh Audio thuyết minh nháp
            with col_btn2:
                st.info("💡 Mẹo: Copy cột 'Lời bình' ở trên dán vào hộp dưới đây để AI đọc thử.")
            
            text_to_read = st.text_area("🎙️ Nhập lời bình (Voiceover) để tạo File Âm thanh:", placeholder="Dán lời thoại vào đây...")
            
            if st.button("🔊 Tạo Giọng Đọc Thuyết Minh (AI Voice)", use_container_width=True):
                if text_to_read.strip():
                    with st.spinner("Đang tổng hợp giọng nói..."):
                        try:
                            # Sử dụng gTTS để đọc tiếng Việt (ngôn ngữ 'vi')
                            tts = gTTS(text=text_to_read, lang='vi', slow=False)
                            audio_buffer = io.BytesIO()
                            tts.write_to_fp(audio_buffer)
                            audio_buffer.seek(0)
                            
                            st.success("Tạo âm thanh thành công! Nghe thử hoặc tải về bên dưới:")
                            st.audio(audio_buffer, format='audio/mp3')
                        except Exception as e:
                            st.error(f"Lỗi tạo giọng nói: {e}. Vui lòng cài đặt thư viện gTTS (pip install gTTS).")
                else:
                    st.warning("Vui lòng nhập nội dung để AI đọc.")
