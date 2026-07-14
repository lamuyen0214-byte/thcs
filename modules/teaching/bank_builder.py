import streamlit as st
import io
import docx
from ai_engine.ai_config import get_api_key
from ai_engine.ai_runner import run_ai_with_fallback

def render_bank_module(api_key=""):
    # Tinh chỉnh CSS
    st.markdown("""
    <style>
    div[data-testid="stAppViewBlockContainer"], .main .block-container { padding-top: 3.5rem !important; padding-bottom: 3rem !important; }
    .stSelectbox label p, .stTextArea label p { color: #0000FF !important; font-weight: 600 !important; font-size: 14px !important; }
    .stMarkdown p strong { color: #FF0000 !important; font-size: 15px !important; }
    .stButton>button { font-weight: bold; border-radius: 6px; }
    </style>
    """, unsafe_allow_html=True)

    # 1. GIAO DIỆN NHẬP LIỆU (CẤU TRÚC MA TRẬN)
    col1, col2, col3 = st.columns([1, 1, 2])
    with col1:
        lop = st.selectbox("Chọn Lớp", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9", "Lớp 10", "Lớp 11", "Lớp 12"], index=0, key="bank_lop")
    with col2:
        mon_hoc = st.selectbox("Chọn Môn", ["KHTN", "Vật Lý", "Hóa Học", "Sinh Học", "Tin học", "Toán", "Ngữ văn", "Lịch sử - Địa lý"], index=0, key="bank_mon")
    with col3:
        loai_de = st.selectbox("Loại đề / Thời gian", ["Kiểm tra 15 phút (10 câu)", "Kiểm tra 1 tiết/Giữa kỳ (40 câu)", "Thi Học kỳ (40-50 câu)", "Đề thi HSG / Chuyên sâu", "Tạo nhóm câu hỏi lẻ"], index=1, key="bank_loai")

    chu_de = st.text_input("Chủ đề / Giới hạn kiến thức", placeholder="Ví dụ: Chương 1: Động học chất điểm, Di truyền học Menđen...")
    
    # Khu vực dán Ma trận đặc tả
    ma_tran = st.text_area("Khung Ma trận & Đặc tả (Tùy chọn nhưng Rất khuyến nghị)", 
                           placeholder="Dán ma trận vào đây. Ví dụ:\n- Nhận biết: 4 câu về khái niệm lực.\n- Thông hiểu: 3 câu tính công cơ học.\n- Vận dụng: 2 câu giải bài toán ròng rọc.\n- Vận dụng cao: 1 câu thực tế...", height=120)

    # 2. KHỞI TẠO AI SINH ĐỀ
    if st.button("TỰ ĐỘNG BIÊN SOẠN ĐỀ THI", type="primary", use_container_width=True):
        if not chu_de.strip() and not ma_tran.strip():
            st.warning("Vui lòng nhập Chủ đề hoặc dán Ma trận để AI biết giới hạn kiến thức!")
            st.stop()
            
        final_key = api_key if api_key else get_api_key()
        
        prompt = f"""
Bạn là Chuyên gia Khảo thí môn {mon_hoc} cấp Trung học.
Hãy biên soạn một {loai_de} dành cho {lop}.
Chủ đề kiến thức: {chu_de}

Cấu trúc Ma trận & Đặc tả giáo viên yêu cầu:
{ma_tran if ma_tran else "Tự động phân bổ theo tỷ lệ 40% Nhận biết - 30% Thông hiểu - 20% Vận dụng - 10% Vận dụng cao."}

YÊU CẦU BẮT BUỘC KHI XUẤT ĐỀ:
1. Soạn đúng số lượng câu hỏi và bám sát tỷ lệ mức độ nhận thức (NB, TH, VD, VDC).
2. TRÌNH BÀY ĐỀ: Phải đánh số thứ tự câu rõ ràng (Câu 1, Câu 2...). Với câu hỏi trắc nghiệm MCQ, các đáp án phải đánh nhãn A, B, C, D trên các dòng riêng biệt.
3. PHẦN ĐÁP ÁN: Tạo một bảng tóm tắt đáp án ở cuối đề thi.
4. PHẦN GIẢI THÍCH: Cung cấp lời giải chi tiết cho các câu Vận dụng và Vận dụng cao.
"""
        with st.spinner("AI đang tìm kiếm ngân hàng câu hỏi và lắp ráp đề thi chuẩn ma trận..."):
            try:
                result = run_ai_with_fallback(prompt=prompt, api_key=final_key, model_mode="flash")
                if result.get("success"):
                    st.session_state['current_exam_bank'] = {"title": chu_de, "content": result.get("text")}
                else:
                    st.error(f"❌ Lỗi AI: {result.get('error')}")
            except Exception as e:
                st.error(f"❌ Lỗi hệ thống: {str(e)}")

    # 3. KẾT QUẢ VÀ XUẤT FILE WORD
    if 'current_exam_bank' in st.session_state:
        st.markdown("---")
        with st.expander("📝 XEM TRƯỚC ĐỀ THI & ĐÁP ÁN", expanded=True):
            st.markdown(st.session_state['current_exam_bank']['content'])
            
            def export_exam_to_word(data):
                doc = docx.Document()
                doc.add_heading(f"ĐỀ KIỂM TRA MÔN {mon_hoc.upper()} - {lop.upper()}", 0)
                if data['title']:
                    doc.add_heading(f"Chủ đề: {data['title']}", 1)
                doc.add_paragraph(data['content'])
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer

            if st.button("📥 Tải Đề thi (Word)", use_container_width=True):
                word_file = export_exam_to_word(st.session_state['current_exam_bank'])
                filename = f"De_Thi_{mon_hoc}_{lop}.docx".replace(" ", "_")
                st.download_button("Xác nhận Tải Đề thi", data=word_file, file_name=filename, use_container_width=True)
