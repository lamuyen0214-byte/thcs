import streamlit as st
import os
import docx
import io
from pypdf import PdfReader
from export.export_word import WordExportEngine
from ai_engine.ai_config import get_api_key
from ai_engine.ai_runner import run_ai_with_fallback

def render_game_module(api_key=""):
    # Tinh chỉnh CSS đồng bộ hệ sinh thái
    st.markdown("""
    <style>
    div[data-testid="stAppViewBlockContainer"], .main .block-container {
        padding-top: 3.5rem !important;
        padding-bottom: 3rem !important;
    }
    .stSelectbox label p, .stTextInput label p, .stNumberInput label p, .stFileUploader label p, .stTextArea label p {
        color: #0000FF !important; 
        font-weight: 600 !important;
        font-size: 14px !important;
    }
    .stMarkdown p strong {
        color: #FF0000 !important; 
        font-size: 15px !important;
    }
    .stButton>button { 
        font-weight: bold; 
        border-radius: 6px;
    }
    </style>
    """, unsafe_allow_html=True)

    # 1. GIAO DIỆN NHẬP LIỆU
    col1, col2, col3 = st.columns([1, 1, 1.5])
    with col1: 
        lop = st.selectbox("Chọn Lớp", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9", "Lớp 10", "Lớp 11", "Lớp 12"], index=0, key="game_lop")
    with col2: 
        mon_hoc = st.selectbox("Chọn Môn", ["KHTN", "Vật Lý", "Hóa Học", "Sinh Học", "Tin học", "Toán", "Ngữ văn", "Lịch sử - Địa lý"], index=0, key="game_mon")
    with col3: 
        loai_tro_choi = st.selectbox("Hình thức Trò chơi", [
            "🧩 Giải Ô chữ (Crossword)", 
            "🎲 Lô tô học tập (Bingo)", 
            "🕵️ Giải mã mật thư (Escape Room)", 
            "⚡ Đấu trường chớp nhoáng (Trivia)"
        ], key="game_type")

    col_ten, col_file = st.columns([2, 1])
    with col_ten: 
        chu_de = st.text_input("Chủ đề / Kiến thức trọng tâm", placeholder="Ví dụ: Cấu tạo tế bào, Các triều đại phong kiến...")
    with col_file: 
        file_tl = st.file_uploader("Tải tài liệu giới hạn (Docx/PDF)", type=['docx', 'pdf'], key="game_file")
    
    yeu_cau = st.text_area("Yêu cầu kịch bản bổ sung", placeholder="Ví dụ: Ô chữ gồm 10 hàng ngang và 1 từ khóa hàng dọc. Câu hỏi mang tính hài hước...", height=70)

    # 2. KHỞI TẠO AI
    if st.button("TỰ ĐỘNG THIẾT KẾ TRÒ CHƠI HỌC TẬP", type="primary", use_container_width=True):
        if not chu_de.strip(): 
            st.warning("Vui lòng nhập Chủ đề / Kiến thức trọng tâm!")
            st.stop()
        
        file_context = ""
        if file_tl:
            try:
                if file_tl.name.endswith('.pdf'):
                    reader = PdfReader(file_tl)
                    file_context = "\n".join([p.extract_text() for p in reader.pages[:10]])
                else:
                    doc = docx.Document(file_tl)
                    file_context = "\n".join([p.text for p in doc.paragraphs])
            except: 
                st.error("Lỗi đọc file tài liệu đính kèm.")

        final_key = api_key if api_key else get_api_key()
        
        # PROMPT CHUYÊN SÂU CHO GAMIFICATION
        prompt = f"""
Bạn là một chuyên gia Gamification (Trò chơi hóa) trong giáo dục. 
Hãy thiết kế một kịch bản trò chơi dạng '{loai_tro_choi}' cho học sinh {lop} môn {mon_hoc}.
Chủ đề: {chu_de}.
Dữ liệu kiến thức tham khảo: {file_context[:3000]}
Yêu cầu kịch bản của giáo viên: {yeu_cau}

Hãy xuất nội dung theo cấu trúc tương ứng với loại trò chơi (Sử dụng định dạng Markdown rõ ràng):
- NẾU LÀ Ô CHỮ: Liệt kê danh sách các Từ khóa đáp án và Câu hỏi gợi ý tương ứng. Chọn sẵn 1 từ khóa hàng dọc (chủ đề chính).
- NẾU LÀ BINGO: Tạo danh sách 16-25 thuật ngữ ngắn gọn (để HS điền vào bảng) và Lời gọi/Câu hỏi mô tả tương ứng (để GV đọc).
- NẾU LÀ GIẢI MÃ MẬT THƯ: Viết một cốt truyện ngắn hấp dẫn (VD: Bị nhốt trong phòng lab...), tạo 4 trạm thử thách (câu hỏi logic/kiến thức), đáp án mỗi trạm ghép lại thành Mật khẩu thoát hiểm cuối cùng.
- NẾU LÀ TRIVIA: 15 câu hỏi trắc nghiệm nhanh, siêu ngắn, có yếu tố gây nhiễu vui nhộn, kèm đáp án.

Yêu cầu chung: Ngôn từ phù hợp lứa tuổi học sinh, kích thích sự tò mò và hứng thú.
"""
        
        with st.spinner(f"AI đang sáng tạo kịch bản {loai_tro_choi}..."):
            try:
                # Sử dụng gemini-1.5-flash để đảm bảo tốc độ và độ ổn định
                result = run_ai_with_fallback(prompt=prompt, api_key=final_key, model_mode="flash")
                if result.get("success"):
                    st.session_state['current_game_data'] = {"title": chu_de, "content": result.get("text")}
                    st.rerun()
                else: 
                    st.error(f"❌ AI từ chối phản hồi. Chi tiết lỗi: {result.get('error')}")
            except Exception as e:
                st.error(f"❌ Lỗi hệ thống: {str(e)}")

    # 3. KẾT QUẢ VÀ XUẤT FILE
    if 'current_game_data' in st.session_state:
        st.markdown("---")
        with st.expander("🎮 XEM TRƯỚC KỊCH BẢN TRÒ CHƠI", expanded=True):
            st.markdown(st.session_state['current_game_data']['content'])
            
            # Khởi tạo hàm xuất Word tại chỗ để tránh lỗi dính form Ma Trận
            def export_game_to_word(data):
                doc = docx.Document()
                doc.add_heading(f"KỊCH BẢN TRÒ CHƠI: {data['title'].upper()}", 0)
                doc.add_paragraph(data['content'])
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer

            if st.button("📥 Tải Kịch bản & Phiếu trò chơi (Word)", use_container_width=True):
                word_file = export_game_to_word(st.session_state['current_game_data'])
                st.download_button("Xác nhận Tải về", data=word_file, file_name=f"Tro_Choi_{chu_de.replace(' ', '_')}.docx", use_container_width=True)
