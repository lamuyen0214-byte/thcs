import streamlit as st
import streamlit.components.v1 as components
import docx
import io
from ai_engine.ai_config import get_api_key
from ai_engine.ai_runner import run_ai_with_fallback

def render_simulation_module(api_key=""):
    # Tinh chỉnh CSS đồng bộ
    st.markdown("""
    <style>
    div[data-testid="stAppViewBlockContainer"], .main .block-container {
        padding-top: 3.5rem !important;
        padding-bottom: 3rem !important;
    }
    .stSelectbox label p, .stTextInput label p, .stNumberInput label p {
        color: #0000FF !important; 
        font-weight: 600 !important;
        font-size: 14px !important;
    }
    .stMarkdown p strong { color: #FF0000 !important; font-size: 15px !important; }
    .stButton>button { font-weight: bold; border-radius: 6px; }
    </style>
    """, unsafe_allow_html=True)

    # DANH SÁCH MÔ PHỎNG PHET MẪU (BẢN TIẾNG VIỆT)
    phet_sims = {
        "Vật lý: Lực và Chuyển động cơ bản": "https://phet.colorado.edu/sims/html/forces-and-motion-basics/latest/forces-and-motion-basics_vi.html",
        "Vật lý: Công viên năng lượng (Trượt ván)": "https://phet.colorado.edu/sims/html/energy-skate-park/latest/energy-skate-park_vi.html",
        "Hóa học: Xây dựng Nguyên tử": "https://phet.colorado.edu/sims/html/build-an-atom/latest/build-an-atom_vi.html",
        "Hóa học: Trạng thái vật chất": "https://phet.colorado.edu/sims/html/states-of-matter/latest/states-of-matter_vi.html",
        "Toán học: Ghép phân số": "https://phet.colorado.edu/sims/html/fraction-matcher/latest/fraction-matcher_vi.html",
        "Tùy chỉnh (Nhập link PhET/HTML5 khác)...": "custom"
    }

    # 1. GIAO DIỆN NHẬP LIỆU
    col1, col2 = st.columns([2, 1])
    with col1:
        chu_de_sim = st.selectbox("Chọn kho Mô phỏng trực quan (PhET Interactive Simulations)", list(phet_sims.keys()))
        
        if chu_de_sim == "Tùy chỉnh (Nhập link PhET/HTML5 khác)...":
            embed_url = st.text_input("Dán đường link Embed (HTML5) vào đây:")
        else:
            embed_url = phet_sims[chu_de_sim]
            
    with col2:
        lop = st.selectbox("Dành cho", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9", "Lớp 10", "Lớp 11", "Lớp 12"], index=2)

    yeu_cau = st.text_input("Yêu cầu AI sinh Phiếu thực hành", placeholder="Ví dụ: Tạo 3 câu hỏi khai thác về sự ma sát...")

    # 2. KHU VỰC HIỂN THỊ MÔ PHỎNG TRỰC QUAN (IFRAME)
    st.markdown("---")
    st.markdown(f"### 🖥️ Màn hình tương tác: {chu_de_sim}")
    if embed_url:
        # Nhúng trực tiếp iframe mô phỏng vào Streamlit
        components.iframe(embed_url, width=800, height=500, scrolling=False)
    else:
        st.info("Vui lòng chọn hoặc nhập link mô phỏng để hiển thị.")

    # 3. AI SINH PHIẾU HƯỚNG DẪN DỰA TRÊN MÔ PHỎNG
    if st.button("TỰ ĐỘNG SOẠN PHIẾU HƯỚNG DẪN KHAI THÁC MÔ PHỎNG NÀY", type="primary"):
        if not embed_url: 
            st.warning("Cần có mô phỏng để soạn phiếu!")
            st.stop()
            
        final_key = api_key if api_key else get_api_key()
        
        prompt = f"""
Bạn là chuyên gia sư phạm KHTN. Giáo viên đang trình chiếu mô phỏng tương tác có chủ đề: '{chu_de_sim}' cho học sinh {lop}.
Yêu cầu của giáo viên: {yeu_cau}

Hãy soạn một "Phiếu học tập & Khám phá mô phỏng" để phát cho học sinh thao tác theo.
Cấu trúc yêu cầu:
1. MỤC TIÊU KHÁM PHÁ (Ngắn gọn)
2. NHIỆM VỤ THỰC HÀNH (Từng bước yêu cầu học sinh kéo/thả/click gì trong mô phỏng).
3. BẢNG GHI CHÚ KẾT QUẢ (Vẽ bảng bằng Markdown để học sinh điền số liệu/hiện tượng quan sát được).
4. CÂU HỎI RÚT RA QUY LUẬT (Khoảng 3-4 câu hỏi phân tích sâu từ hiện tượng vừa xem).

Nội dung phải sinh động, kích thích tư duy khoa học!
"""
        with st.spinner("AI đang "xem" mô phỏng và soạn phiếu bài tập..."):
            try:
                result = run_ai_with_fallback(prompt=prompt, api_key=final_key, model_mode="flash")
                if result.get("success"):
                    st.session_state['current_sim_worksheet'] = {"title": chu_de_sim, "content": result.get("text")}
                else: 
                    st.error(f"❌ Lỗi AI: {result.get('error')}")
            except Exception as e:
                st.error(f"❌ Lỗi hệ thống: {str(e)}")

    # 4. KẾT QUẢ PHIẾU HƯỚNG DẪN VÀ XUẤT WORD
    if 'current_sim_worksheet' in st.session_state:
        st.markdown("---")
        with st.expander("📝 XEM PHIẾU HỌC TẬP KHAI THÁC MÔ PHỎNG", expanded=True):
            st.markdown(st.session_state['current_sim_worksheet']['content'])
            
            def export_sim_to_word(data):
                doc = docx.Document()
                doc.add_heading(f"PHIẾU HỌC TẬP MÔ PHỎNG: {data['title']}", 0)
                doc.add_paragraph(data['content'])
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer

            if st.button("📥 Tải Phiếu học tập (Word)", use_container_width=True):
                word_file = export_sim_to_word(st.session_state['current_sim_worksheet'])
                st.download_button("Xác nhận Tải về", data=word_file, file_name="Phieu_Mo_Phong.docx", use_container_width=True)
