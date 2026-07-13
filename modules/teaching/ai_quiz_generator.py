import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
import time
from pypdf import PdfReader

# 1. HÀM KHỞI TẠO AI AN TOÀN
def get_stable_model(api_key):
    if not api_key:
        st.error("API Key chưa được cấu hình trong Sidebar!")
        return None
    try:
        genai.configure(api_key=api_key)
        # Sử dụng các model ổn định
        priority = ["gemini-1.5-flash", "gemini-2.0-flash-lite", "gemini-flash-latest"]
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        
        for name in priority:
            if name in available_models: 
                return genai.GenerativeModel(name)
        
        # Fallback nếu không thấy model ưu tiên
        return genai.GenerativeModel("gemini-1.5-flash")
    except Exception as e:
        st.error(f"Lỗi khởi tạo mô hình: {e}")
        return None

# 2. HÀM GIAO DIỆN CHÍNH
def render_quiz_generator(api_key):
    st.subheader("🎯 Trình Tạo Đề Kiểm Tra Tự Động")
    
    # Khu vực đầu vào
    uploaded_file = st.file_uploader("Tải tài liệu tham khảo (PDF, DOCX):", type=['pdf', 'docx', 'txt'])
    text_content = st.text_area("Nội dung bài giảng/Yêu cầu bổ sung (VD: 8 câu MCQ, 1 Đ/S...):", height=200)
    
    col1, col2 = st.columns(2)
    with col1:
        num = st.number_input("Số câu:", min_value=1, value=5)
        q_type = st.selectbox("Dạng bài:", ["Trắc nghiệm", "Tự luận", "Trắc nghiệm kết hợp tự luận"])
    with col2:
        use_source = st.checkbox("Bám sát tài liệu", value=True)
        include_digital = st.checkbox("Tích hợp Năng lực số", value=True)
    
    if st.button("🚀 Tạo đề ngay"):
        # Bước kiểm tra model
        model = get_stable_model(api_key)
        if model is None:
            st.stop() # Dừng tại đây nếu không có model
            
        # Xử lý nội dung tài liệu
        combined = text_content
        if use_source and uploaded_file:
            try:
                if uploaded_file.type == "application/pdf": 
                    combined += "\n" + "\n".join([p.extract_text() for p in PdfReader(uploaded_file).pages])
                elif "word" in uploaded_file.type or uploaded_file.name.endswith('.docx'): 
                    doc = Document(uploaded_file)
                    combined += "\n" + "\n".join([p.text for p in doc.paragraphs])
            except Exception as e:
                st.error(f"Lỗi đọc file: {e}")
        
        if not combined: 
            st.warning("Vui lòng nhập nội dung bài giảng hoặc tải tài liệu lên!")
        else:
            with st.spinner("AI đang soạn đề..."):
                prompt = f"Soạn {num} câu hỏi {q_type}. {'Lồng ghép Năng lực số.' if include_digital else ''} Dựa trên: {combined}."
                
                # Vòng lặp thử lại nếu lỗi quá tải
                success = False
                for i in range(3):
                    try:
                        res = model.generate_content(prompt)
                        st.session_state.current_quiz = res.text
                        success = True
                        st.rerun()
                        break
                    except Exception as e:
                        if "429" in str(e) and i < 2:
                            time.sleep(20)
                        else:
                            st.error(f"Lỗi AI: {e}")
                            break

    # Hiển thị kết quả
    if "current_quiz" in st.session_state:
        st.markdown("---")
        st.text_area("Kết quả:", st.session_state.current_quiz, height=300)
        
        # Xuất file Word
        doc = Document()
        doc.add_paragraph(st.session_state.current_quiz)
        bio = BytesIO()
        doc.save(bio)
        st.download_button("📥 Tải bộ đề (.docx)", data=bio.getvalue(), file_name="De_kiem_tra.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
