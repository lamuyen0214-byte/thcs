import streamlit as st
import io
import docx
from ai_engine.ai_config import get_api_key
from ai_engine.ai_runner import run_ai_with_fallback

def render_interaction_module(api_key=""):
    # Tinh chỉnh CSS
    st.markdown("""
    <style>
    div[data-testid="stAppViewBlockContainer"], .main .block-container { padding-top: 3.5rem !important; padding-bottom: 3rem !important; }
    .stSelectbox label p, .stTextInput label p, .stTextArea label p { color: #0000FF !important; font-weight: 600 !important; font-size: 14px !important; }
    .stMarkdown p strong { color: #FF0000 !important; font-size: 15px !important; }
    .stButton>button { font-weight: bold; border-radius: 6px; }
    </style>
    """, unsafe_allow_html=True)

    # 1. GIAO DIỆN THIẾT LẬP HOẠT ĐỘNG
    col1, col2 = st.columns([1, 1])
    with col1:
        loai_tuong_tac = st.selectbox("🎭 Hình thức Tương tác", [
            "Đóng vai Nhân vật (Roleplay)", 
            "Tình huống Tranh biện (Debate)", 
            "Giải quyết Vấn đề Thực tế (Case Study)"
        ])
    with col2:
        mon_hoc = st.selectbox("📚 Môn học & Đối tượng", [
            "KHTN - Cấp THCS", "Vật Lý", "Hóa Học", "Sinh Học", 
            "Lịch sử - Địa lý", "Ngữ Văn", "GDCD / Hoạt động Trải nghiệm"
        ])
        
    chu_de = st.text_input("🎯 Chủ đề / Nhân vật cốt lõi", placeholder="Ví dụ: Đóng vai Các-mác, Tranh biện về Năng lượng hạt nhân, Ô nhiễm môi trường nước...")
    yeu_cau = st.text_area("📝 Yêu cầu chi tiết thêm (Tùy chọn)", placeholder="Ví dụ: Lồng ghép yếu tố hài hước, chia thành 2 phe ủng hộ và phản đối rõ ràng...", height=80)

    # 2. KHỞI TẠO AI SÁNG TẠO KỊCH BẢN
    if st.button("TẠO KỊCH BẢN HOẠT ĐỘNG TƯƠNG TÁC", type="primary", use_container_width=True):
        if not chu_de.strip():
            st.warning("Vui lòng nhập Chủ đề hoặc Nhân vật cốt lõi!")
            st.stop()
            
        final_key = api_key if api_key else get_api_key()
        
        # PROMPT THÍCH ỨNG THEO LOẠI TƯƠNG TÁC
        if "Đóng vai" in loai_tuong_tac:
            prompt = f"""
Bạn là chuyên gia Sư phạm chuyên thiết kế hoạt động Roleplay (Đóng vai) cho học sinh {mon_hoc}.
Chủ đề/Nhân vật: {chu_de}
Yêu cầu thêm: {yeu_cau}

Hãy viết một 'Kịch bản Đóng vai' gồm:
1. BỐI CẢNH CÂU CHUYỆN: Ngắn gọn, hấp dẫn để lôi cuốn học sinh.
2. HỒ SƠ NHÂN VẬT: Danh sách các nhân vật cần có, tính cách và mục tiêu của từng người trong vở kịch.
3. KỊCH BẢN MỞ ĐẦU: Các câu thoại mồi (Starter dialogs) để học sinh bắt đầu nhập vai.
4. CÂU HỎI TƯƠNG TÁC LỚP: 3 câu hỏi để giáo viên hỏi những học sinh ngồi dưới đang xem kịch.
"""
        elif "Tranh biện" in loai_tuong_tac:
             prompt = f"""
Bạn là chuyên gia Sư phạm chuyên tổ chức Debate (Tranh biện) cho học sinh {mon_hoc}.
Chủ đề tranh biện: {chu_de}
Yêu cầu thêm: {yeu_cau}

Hãy viết một 'Kế hoạch Tổ chức Tranh biện' gồm:
1. TUYÊN BỐ KIẾN NGHỊ (Motion): Một câu khẳng định rõ ràng để làm đề tài tranh biện.
2. LẬP LUẬN PHE ỦNG HỘ (Affirmative): Gợi ý 3 luận điểm chính để nhóm ủng hộ khai thác.
3. LẬP LUẬN PHE PHẢN ĐỐI (Negative): Gợi ý 3 luận điểm chính để nhóm phản đối khai thác.
4. CÂU HỎI PHẢN BIỆN CHÉO: Gợi ý 2 câu hỏi hóc búa để hai đội "tấn công" nhau.
5. LUẬT CHƠI & TIÊU CHÍ CHẤM ĐIỂM: Ngắn gọn, dễ hiểu cho học sinh THCS/THPT.
"""
        else:
             prompt = f"""
Bạn là chuyên gia Sư phạm chuyên thiết kế Case Study (Tình huống thực tế) cho học sinh {mon_hoc}.
Chủ đề: {chu_de}
Yêu cầu thêm: {yeu_cau}

Hãy viết một 'Kịch bản Giải quyết Vấn đề' gồm:
1. CÂU CHUYỆN TÌNH HUỐNG: Một tình huống thực tế, có tính tiến thoái lưỡng nan (dilemma) hoặc cần vận dụng kiến thức bài học để giải quyết.
2. NHIỆM VỤ CỦA HỌC SINH: Rõ ràng, yêu cầu tạo ra sản phẩm hoặc đưa ra quyết định gì.
3. GỢI Ý GIẢI QUYẾT: Dành cho giáo viên (đáp án mở).
4. PHIẾU BÀI TẬP NHÓM: Khung sườn để học sinh điền ý tưởng giải quyết vào.
"""

        with st.spinner("AI đang thiết kế kịch bản hoạt động tương tác..."):
            try:
                result = run_ai_with_fallback(prompt=prompt, api_key=final_key, model_mode="flash")
                if result.get("success"):
                    st.session_state['current_interaction'] = {"title": chu_de, "content": result.get("text")}
                else:
                    st.error(f"❌ Lỗi AI: {result.get('error')}")
            except Exception as e:
                st.error(f"❌ Lỗi hệ thống: {str(e)}")

    # 3. KẾT QUẢ VÀ XUẤT FILE WORD
    if 'current_interaction' in st.session_state:
        st.markdown("---")
        with st.expander("🎭 XEM TRƯỚC KỊCH BẢN TƯƠNG TÁC", expanded=True):
            st.markdown(st.session_state['current_interaction']['content'])
            
            def export_interaction_to_word(data):
                doc = docx.Document()
                doc.add_heading(f"HOẠT ĐỘNG TƯƠNG TÁC: {data['title'].upper()}", 0)
                doc.add_paragraph(data['content'])
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer

            if st.button("📥 Tải Kịch bản Hoạt động (Word)", use_container_width=True):
                word_file = export_interaction_to_word(st.session_state['current_interaction'])
                st.download_button("Xác nhận Tải về", data=word_file, file_name=f"Kich_Ban_Tuong_Tac.docx", use_container_width=True)
