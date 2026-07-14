import streamlit as st
import os
import docx
import io
from pypdf import PdfReader
from ai_engine.ai_config import get_api_key
from ai_engine.ai_runner import run_ai_with_fallback

def render_hoc_lieu_module(api_key=""):
    # Tinh chỉnh CSS đồng bộ
    st.markdown("""
    <style>
    div[data-testid="stAppViewBlockContainer"], .main .block-container {
        padding-top: 3.5rem !important;
        padding-bottom: 3rem !important;
    }
    .stSelectbox label p, .stTextInput label p, .stNumberInput label p, .stFileUploader label p, .stTextArea label p {
        color: #0000FF !important; 
        font-weight: 600 !important;
        font-size: 14px !important;
    }
    .stMarkdown p strong {
        color: #FF0000 !important; 
        font-size: 15px !important;
    }
    .stButton>button { 
        font-weight: bold; 
        border-radius: 6px;
    }
    </style>
    """, unsafe_allow_html=True)

    # 1. GIAO DIỆN NHẬP LIỆU
    col1, col2, col3 = st.columns([1, 1, 2])
    with col1: 
        lop = st.selectbox("Chọn Lớp", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9", "Lớp 10", "Lớp 11", "Lớp 12"], index=0, key="hl_lop")
    with col2: 
        mon_hoc = st.selectbox("Chọn Môn", ["KHTN", "Vật Lý", "Hóa Học", "Sinh Học", "Tin học", "Toán", "Ngữ văn", "Lịch sử - Địa lý"], index=0, key="hl_mon")
    with col3: 
        loai_hoc_lieu = st.selectbox("Loại Học liệu cần tạo", [
            "📚 Tóm tắt lý thuyết (Đề cương)", 
            "📝 Phiếu bài tập (Worksheet)", 
            "📇 Thẻ ghi nhớ (Flashcard)", 
            "🧠 Cấu trúc Sơ đồ tư duy (Mindmap)"
        ], key="hl_loai")

    col_ten, col_file = st.columns([2, 1])
    with col_ten: 
        chu_de = st.text_input("Chủ đề / Tên bài học", placeholder="Ví dụ: Định luật Newton, Cấu tạo tế bào...")
    with col_file: 
        file_tl = st.file_uploader("Tải tài liệu tham khảo (Docx/PDF)", type=['docx', 'pdf'], key="hl_file")
    
    yeu_cau = st.text_area("Yêu cầu chi tiết", placeholder="Ví dụ: Tóm tắt ngắn gọn dưới 2 trang, Phiếu bài tập chia thành 3 mức độ (Dễ, Trung bình, Khó)...", height=70)

    # 2. KHỞI TẠO AI
    if st.button("TỰ ĐỘNG BIÊN SOẠN HỌC LIỆU", type="primary", use_container_width=True):
        if not chu_de.strip(): 
            st.warning("Vui lòng nhập Chủ đề / Tên bài học!")
            st.stop()
        
        file_context = ""
        if file_tl:
            try:
                if file_tl.name.endswith('.pdf'):
                    reader = PdfReader(file_tl)
                    file_context = "\n".join([p.extract_text() for p in reader.pages[:10]])
                else:
                    doc = docx.Document(file_tl)
                    file_context = "\n".join([p.text for p in doc.paragraphs])
            except: 
                st.error("Lỗi đọc file tài liệu đính kèm.")

        final_key = api_key if api_key else get_api_key()
        
        # PROMPT CHUYÊN SÂU TẠO HỌC LIỆU
        prompt = f"""
Bạn là một chuyên gia thiết kế học liệu giáo dục xuất sắc. Hãy biên soạn một '{loai_hoc_lieu}' cho học sinh {lop} môn {mon_hoc}.
Chủ đề: {chu_de}.
Dữ liệu kiến thức tham khảo (nếu có): {file_context[:3500]}
Yêu cầu bổ sung của giáo viên: {yeu_cau}

Yêu cầu định dạng cấu trúc (Sử dụng Markdown rõ ràng):
- Nếu là "Tóm tắt lý thuyết": Phân chia các đề mục rõ ràng (I, II, 1, 2), bôi đậm từ khóa quan trọng.
- Nếu là "Phiếu bài tập": Có phần tóm tắt công thức/lý thuyết ngắn, sau đó là hệ thống câu hỏi phân loại từ dễ đến khó. Phải cung cấp ĐÁP ÁN ở cuối phiếu.
- Nếu là "Flashcard": Cấu trúc dưới dạng bảng. Cột 1 (Mặt trước: Câu hỏi/Thuật ngữ) - Cột 2 (Mặt sau: Đáp án/Định nghĩa ngắn gọn).
- Nếu là "Sơ đồ tư duy": Trình bày theo dạng phân cấp nhánh cây (Ví dụ: - Nhánh chính 1 \n  - Nhánh phụ 1.1) để giáo viên dễ dàng vẽ lại lên bảng hoặc copy vào phần mềm Mindmap.

Hãy xuất nội dung thật khoa học và chuyên nghiệp.
"""
        
        with st.spinner(f"AI đang biên soạn {loai_hoc_lieu}..."):
            try:
                result = run_ai_with_fallback(prompt=prompt, api_key=final_key, model_mode="flash")
                if result.get("success"):
                    st.session_state['current_hoc_lieu'] = {"title": chu_de, "content": result.get("text")}
                    st.rerun()
                else: 
                    st.error(f"❌ AI từ chối phản hồi: {result.get('error')}")
            except Exception as e:
                st.error(f"❌ Lỗi hệ thống: {str(e)}")

    # 3. KẾT QUẢ VÀ XUẤT FILE WORD
    if 'current_hoc_lieu' in st.session_state:
        st.markdown("---")
        with st.expander("📖 XEM TRƯỚC HỌC LIỆU", expanded=True):
            st.markdown(st.session_state['current_hoc_lieu']['content'])
            
            def export_hoclieu_to_word(data):
                doc = docx.Document()
                doc.add_heading(f"HỌC LIỆU: {data['title'].upper()}", 0)
                doc.add_paragraph(data['content'])
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer

            if st.button("📥 Tải Học liệu về máy (Word)", use_container_width=True):
                word_file = export_hoclieu_to_word(st.session_state['current_hoc_lieu'])
                st.download_button("Xác nhận Tải file", data=word_file, file_name=f"Hoc_Lieu_{chu_de.replace(' ', '_')}.docx", use_container_width=True)
