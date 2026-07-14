import streamlit as st
import pandas as pd
import io
import docx
from ai_engine.ai_config import get_api_key
from ai_engine.ai_runner import run_ai_with_fallback

def load_and_clean_data(file):
    """Hàm xử lý thông minh để dọn rác file Excel từ SMAS, vnEdu"""
    try:
        if file.name.endswith('.csv'):
            return pd.read_csv(file)
            
        # Đọc file không dùng dòng đầu làm tiêu đề để lấy toàn bộ dữ liệu
        df_raw = pd.read_excel(file, header=None)
        
        # Tìm dòng có chứa chữ "Họ và tên" hoặc "Họ tên" (Giới hạn tìm trong 20 dòng đầu)
        header_idx = -1
        for i in range(min(20, len(df_raw))):
            row_values = df_raw.iloc[i].astype(str).str.lower().values
            if any('họ và tên' in val or 'họ tên' in val for val in row_values):
                header_idx = i
                break
                
        if header_idx != -1:
            # PHÁT HIỆN CẤU TRÚC SMAS/VNEDU
            # Lấy dòng tiêu đề chính (kéo dài các ô bị merge bằng ffill)
            header_row_1 = df_raw.iloc[header_idx].ffill()
            
            # Lấy dòng tiêu đề phụ (Cột 1, 2, 3, 4 của ĐĐG TX)
            if header_idx + 1 < len(df_raw):
                header_row_2 = df_raw.iloc[header_idx + 1].fillna('')
            else:
                header_row_2 = [''] * len(header_row_1)
            
            # Ghép tiêu đề lại cho phẳng (Ví dụ: "ĐĐG TX (1)")
            new_headers = []
            for h1, h2 in zip(header_row_1, header_row_2):
                h1_str = str(h1).strip() if pd.notna(h1) else ""
                h2_str = str(h2).strip() if pd.notna(h2) and str(h2).strip() not in ['', 'nan'] else ""
                
                if h2_str:
                    new_headers.append(f"{h1_str} ({h2_str})")
                else:
                    new_headers.append(h1_str)
                    
            # Cắt bỏ toàn bộ rác ở trên, chỉ lấy phần dữ liệu học sinh
            df_clean = df_raw.iloc[header_idx + 2:].copy()
            df_clean.columns = new_headers
            
            # Xóa các cột trống không có tên
            df_clean = df_clean.loc[:, [c for c in df_clean.columns if str(c).lower() != 'nan' and c != '']]
            
            # Xóa các dòng rỗng (Không có dữ liệu ở cột Họ và tên)
            name_col = [c for c in df_clean.columns if 'họ và tên' in c.lower() or 'họ tên' in c.lower()][0]
            df_clean = df_clean.dropna(subset=[name_col])
            
            # Tự động ép kiểu các cột Điểm về dạng số học để thống kê
            for col in df_clean.columns:
                if any(x in col for x in ['ĐĐG', 'TBM', 'Điểm', 'TX', 'GK', 'CK']):
                    df_clean[col] = pd.to_numeric(df_clean[col], errors='coerce')
            
            return df_clean
            
        else:
            # Nếu là file Excel chuẩn mực bình thường, đọc như cũ
            return pd.read_excel(file)
            
    except Exception as e:
        raise Exception(f"Lỗi khi dọn dẹp dữ liệu: {e}")


def render_analytics_module(api_key=""):
    # Tinh chỉnh CSS
    st.markdown("""
    <style>
    div[data-testid="stAppViewBlockContainer"], .main .block-container { padding-top: 3.5rem !important; padding-bottom: 3rem !important; }
    .stSelectbox label p, .stFileUploader label p, .stTextArea label p { color: #0000FF !important; font-weight: 600 !important; font-size: 14px !important; }
    .stMarkdown p strong { color: #FF0000 !important; font-size: 15px !important; }
    .stButton>button { font-weight: bold; border-radius: 6px; }
    </style>
    """, unsafe_allow_html=True)

    # 1. GIAO DIỆN NHẬP LIỆU ĐA CHIỀU
    col_mode, col_req = st.columns([1, 1])
    with col_mode:
        che_do = st.selectbox("🎯 Chọn Chế độ Phân tích", [
            "1. Phân tích Tổng quan Phổ điểm & Xếp loại",
            "2. Phân tích Chất lượng Câu hỏi (Item Analysis)",
            "3. Phân tích Đánh giá Năng lực / Rubric Dự án"
        ])
        file_diem = st.file_uploader("📥 Tải Bảng điểm (Hỗ trợ file xuất từ SMAS, vnEdu, Excel, CSV)", type=['xlsx', 'xls', 'csv'])
    
    with col_req:
        yeu_cau = st.text_area("Yêu cầu trọng tâm cho AI (Tùy chọn)", placeholder="Ví dụ: Chỉ ra 5 học sinh yếu nhất cần phụ đạo gấp, phân tích xem học sinh hổng kiến thức ở câu nào nhiều nhất...", height=115)

    # 2. XỬ LÝ DASHBOARD & AI PHÂN TÍCH
    if file_diem is not None:
        try:
            # SỬ DỤNG HÀM LỌC SMAS THÔNG MINH
            df = load_and_clean_data(file_diem)
            
            st.markdown("---")
            st.markdown("### 📊 Dashboard Thống kê Nhanh")
            
            # Giao diện Tabs cho Dashboard
            tab_data, tab_stats = st.tabs(["1️⃣ Dữ liệu Đã làm sạch", "2️⃣ Thống kê độ phân tán"])
            with tab_data:
                st.dataframe(df.head(15), use_container_width=True) # Hiển thị 15 dòng đầu siêu sạch đẹp
            with tab_stats:
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    st.write(df[numeric_cols].describe())
                else:
                    st.info("Hệ thống không tìm thấy cột số nào để thống kê (Có thể bảng điểm chỉ chứa chữ).")

            if st.button("🤖 TIẾN HÀNH PHÂN TÍCH CHUYÊN SÂU", type="primary", use_container_width=True):
                data_str = df.head(150).to_csv(index=False)
                final_key = api_key if api_key else get_api_key()
                
                if "1" in che_do:
                    sys_role = "Chuyên gia Thống kê Giáo dục"
                    focus = "- Tỷ lệ phân bổ điểm (Giỏi/Khá/TB/Yếu).\n- Điểm mạnh chung của toàn lớp.\n- Phân nhóm học sinh: Nhóm nòng cốt và Nhóm rủi ro cần phụ đạo."
                elif "2" in che_do:
                    sys_role = "Chuyên gia Khảo thí và Đo lường"
                    focus = "- Nhận diện các câu hỏi/phần thi có tỷ lệ làm sai cao nhất.\n- Chẩn đoán nguyên nhân hổng kiến thức từ những câu sai đó.\n- Đề xuất điều chỉnh ma trận đề thi cho lần sau."
                else:
                    sys_role = "Chuyên gia Đánh giá Năng lực"
                    focus = "- Sự đồng đều giữa các tiêu chí (VD: Điểm thuyết trình vs Điểm sản phẩm).\n- Đánh giá kỹ năng thực hành/làm việc nhóm.\n- Nhận xét định tính rút ra từ dữ liệu."

                prompt = f"""
Bạn là {sys_role} hàng đầu. Dưới đây là bảng điểm đã được làm sạch của học sinh (định dạng CSV):
{data_str}

Yêu cầu cụ thể của giáo viên: {yeu_cau}

Dựa vào bảng số liệu trên, hãy viết một "BÁO CÁO SƯ PHẠM" bám sát các trọng tâm sau:
{focus}

Yêu cầu trình bày:
- Dùng định dạng Markdown rõ ràng, chia mục khoa học.
- Dùng bảng biểu Markdown nếu cần so sánh.
- TUYỆT ĐỐI trung thực với số liệu, không bịa đặt tên học sinh hoặc điểm số không có trong file.
"""
                with st.spinner(f"AI đang phân tích dữ liệu SMAS và kích hoạt {che_do}..."):
                    result = run_ai_with_fallback(prompt=prompt, api_key=final_key, model_mode="flash")
                    if result.get("success"):
                        st.session_state['current_analytics'] = result.get("text")
                    else:
                        st.error(f"❌ Lỗi AI: {result.get('error')}")

        except Exception as e:
            st.error(f"❌ Lỗi xử lý file: {str(e)}")

    # 3. KẾT QUẢ VÀ XUẤT FILE WORD
    if 'current_analytics' in st.session_state:
        st.markdown("---")
        with st.expander("📑 BÁO CÁO PHÂN TÍCH CHUYÊN SÂU SƯ PHẠM", expanded=True):
            st.markdown(st.session_state['current_analytics'])
            
            def export_analytics_to_word(content):
                doc = docx.Document()
                doc.add_heading(f"BÁO CÁO ĐÁNH GIÁ SƯ PHẠM", 0)
                doc.add_paragraph(content)
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer

            if st.button("📥 Tải Báo cáo Phân tích (Word)", use_container_width=True):
                word_file = export_analytics_to_word(st.session_state['current_analytics'])
                st.download_button("Xác nhận Tải Báo cáo", data=word_file, file_name="Bao_Cao_Spham_AI.docx", use_container_width=True)
