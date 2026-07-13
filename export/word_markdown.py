"""
Module: export/word_markdown.py
Nhiệm vụ: Chuyển đổi Markdown và Toán học (LaTeX) sang định dạng Microsoft Word nguyên bản.
Kiến trúc: Ép kiểu Font tuyệt đối, xử lý phân số đệ quy chống gãy cấu trúc.
"""
import re
from docx.shared import Pt, RGBColor

class MarkdownParser:
    @staticmethod
    def _force_font(run, is_bold=False, is_italic=False):
        """Hàm bạo chúa: Ép buộc mọi đoạn text phải theo đúng chuẩn GDPT"""
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        if is_bold: run.bold = True
        if is_italic: run.italic = True

    @staticmethod
    def clean_math(text: str) -> str:
        """Bộ lọc chuyên sâu: Chuyển đổi Toán học & bóc tách ngoặc lồng nhau"""
        text = text.replace("$", "").replace("\\[", "").replace("\\]", "")
        
        # 1. Ánh xạ trực tiếp các ký tự Toán học (Đã bổ sung \cdot, \times, \Rightarrow...)
        exact_map = {
            r'\neq': '≠', r'\pm': '±', r'\triangle': '△', r'\cdot': '·', r'\times': '×',
            r'\mathbb{R}': 'ℝ', r'mathbb{R}': 'ℝ', r'\mathbb{Z}': 'ℤ', r'\mathbb{N}': 'ℕ',
            r'\in': '∈', r'\le': '≤', r'\ge': '≥', r'\leq': '≤', r'\geq': '≥',
            r'\pi': 'π', r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ',
            r'\sqrt': '√', r'\rightarrow': '→', r'\leftrightarrow': '↔',
            r'\Rightarrow': '⇒', r'\Leftrightarrow': '⇔',
            r'\infty': '∞', '^2': '²', '^3': '³', r'\circ': '°'
        }
        for tex, uni in exact_map.items():
            text = text.replace(tex, uni)
            
        # 2. Ánh xạ dùng ranh giới từ (Tránh làm hỏng Tiếng Việt)
        safe_regex = {
            r'\bneq\b': '≠', r'\bpm\b': '±', r'\btriangle\b': '△',
            r'\ble\b': '≤', r'\bge\b': '≥', r'\bpi\b': 'π',
            r'\balpha\b': 'α', r'\bbeta\b': 'β', r'\binfty\b': '∞',
            r'\bx in\b(?=\s*ℝ)': 'x ∈', r'\bx in\b(?=\s*math)': 'x ∈'
        }
        for pat, uni in safe_regex.items():
            text = re.sub(pat, uni, text)
            
        # 3. Vòng lặp đệ quy bóc tách phân số \frac{A}{B} (Giải quyết triệt để lỗi ngoặc lồng nhau)
        # Quét liên tục cho đến khi không còn \frac nào trong chuỗi
        frac_pattern = r'\\?frac\{((?:[^{}]|\{[^{}]*\})*)\}\{((?:[^{}]|\{[^{}]*\})*)\}'
        while re.search(frac_pattern, text):
            text = re.sub(frac_pattern, r'(\1)/(\2)', text)
            
        return text

    @staticmethod
    def _parse_inline_formatting(paragraph, text: str):
        """Cắt chuỗi thông minh để xử lý In đậm (**), In nghiêng (*)"""
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if not part: continue
            
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                MarkdownParser._force_font(run, is_bold=True)
            elif part.startswith('*') and part.endswith('*'):
                # Kiểm tra tránh bắt nhầm dấu * đơn lẻ của phép nhân
                if len(part) > 2 and part[1] != ' ' and part[-2] != ' ':
                    run = paragraph.add_run(part[1:-1])
                    MarkdownParser._force_font(run, is_italic=True)
                else:
                    run = paragraph.add_run(part)
                    MarkdownParser._force_font(run)
            else:
                run = paragraph.add_run(part)
                MarkdownParser._force_font(run)

    @staticmethod
    def _build_word_table(doc, table_rows_data):
        """Kiến tạo Bảng Word tự động - Lọc rác và cân bằng cột"""
        if not table_rows_data: return

        # Làm sạch mảng: Loại bỏ dòng kẻ ngang Markdown (vd: |---|---|)
        valid_rows = []
        for r in table_rows_data:
            if set(r.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
                continue 
            valid_rows.append(r)
            
        if not valid_rows: return

        # Tính số cột lớn nhất để tránh lỗi Out Of Range khi AI tạo bảng thiếu ô
        num_cols = max(len([c for c in row.split('|') if c.strip() or c == '']) - 2 for row in valid_rows)
        if num_cols <= 0: return

        table = doc.add_table(rows=len(valid_rows), cols=num_cols)
        table.style = 'Table Grid'

        for r_idx, row_text in enumerate(valid_rows):
            cells = row_text.split('|')[1:-1] 
            for c_idx in range(min(num_cols, len(cells))):
                cell_text = cells[c_idx].strip()
                cell_p = table.cell(r_idx, c_idx).paragraphs[0]
                MarkdownParser._parse_inline_formatting(cell_p, MarkdownParser.clean_math(cell_text))
                
                # In đậm hàng tiêu đề đầu tiên
                if r_idx == 0:
                    for run in cell_p.runs: run.bold = True

    @staticmethod
    def parse(doc, markdown_text: str):
        """Hàm điều phối trung tâm: Đọc từng dòng và vẽ ra Word"""
        lines = markdown_text.split('\n')
        in_table = False
        table_rows = []

        for line in lines:
            raw_line = line.strip()
            
            # Xóa sạch các đường kẻ ngang rác của AI
            if raw_line == '---' or raw_line == '***':
                continue

            # BẮT BẢNG (TABLE)
            if raw_line.startswith('|') and raw_line.endswith('|'):
                in_table = True
                table_rows.append(raw_line)
                continue
            else:
                if in_table:
                    MarkdownParser._build_word_table(doc, table_rows)
                    in_table = False
                    table_rows = []
                    doc.add_paragraph() 
            
            if not raw_line: continue

            # Dọn dẹp Toán học trước khi xét định dạng
            line_clean = MarkdownParser.clean_math(raw_line)

            # XỬ LÝ TIÊU ĐỀ (HEADING)
            heading_match = re.match(r'^(#{1,6})\s+(.*)', line_clean)
            if heading_match:
                # Không dùng level của Word để tránh bị đổi Font, tự mô phỏng Heading
                h = doc.add_paragraph()
                MarkdownParser._parse_inline_formatting(h, heading_match.group(2))
                for run in h.runs: 
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.size = Pt(14) # Heading to hơn 1 chút
                    run.bold = True
                continue

            # XỬ LÝ DANH SÁCH BULLET (Dấu chấm tròn)
            list_match = re.match(r'^[\*\-]\s+(.*)', line_clean)
            if list_match:
                p = doc.add_paragraph(style='List Bullet')
                MarkdownParser._parse_inline_formatting(p, list_match.group(1))
                continue
                
            # XỬ LÝ DANH SÁCH NUMBER (Đánh số 1. 2. 3.)
            num_match = re.match(r'^(\d+\.)\s+(.*)', line_clean)
            if num_match:
                p = doc.add_paragraph(style='List Number')
                MarkdownParser._parse_inline_formatting(p, num_match.group(2))
                continue

            # VĂN BẢN THƯỜNG
            p = doc.add_paragraph()
            MarkdownParser._parse_inline_formatting(p, line_clean)

        # Chốt hạ bảng nếu nó nằm ở cuối cùng
        if in_table:
            MarkdownParser._build_word_table(doc, table_rows)
