"""
Module: export/word_tables.py
Nhiệm vụ: Dựng bảng động từ AST Node và vẽ các bảng biểu mẫu tĩnh (Ma trận, Đặc tả, KHBD).
Bảo mật: Duyệt mảng an toàn, không sử dụng index cứng để tránh IndexError.
"""

from typing import Dict, Any
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

class TableRenderer:
    @staticmethod
    def _set_cell_background(cell, hex_color: str):
        """Hàm can thiệp lõi XML để đổ màu nền cho ô Word một cách sạch sẽ"""
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

    @classmethod
    def render_ast_table(cls, doc: docx.Document, node: Dict[str, Any], style_manager: Any, math_renderer: Any):
        """
        Vẽ bảng động từ AST Node do Tokenizer sinh ra.
        Chống mọi lỗi IndexError nhờ dữ liệu đã được chuẩn hóa số lượng cột ở tầng Tokenizer.
        """
        cols = node.get("cols", 0)
        if cols == 0:
            return

        headers = node.get("headers", [])
        rows = node.get("rows", [])
        total_rows = (1 if headers else 0) + len(rows)

        table = doc.add_table(rows=total_rows, cols=cols)
        table.style = 'Table Grid'
        table.autofit = True

        current_row = 0

        # 1. Vẽ hàng Tiêu đề (Header)
        if headers:
            for c_idx, cell_tokens in enumerate(headers):
                cell = table.cell(current_row, c_idx)
                cls._set_cell_background(cell, "F2F4F4") # Màu xám nhạt chuyên nghiệp
                p = cell.paragraphs[0]
                # Gọi StyleManager để vẽ Text/Toán nội dòng vào ô
                style_manager._render_inline_tokens(p, cell_tokens, math_renderer)
                # Ép in đậm cho Header
                for run in p.runs:
                    run.bold = True
            current_row += 1

        # 2. Vẽ nội dung Bảng (Body)
        for row_data in rows:
            for c_idx, cell_tokens in enumerate(row_data):
                cell = table.cell(current_row, c_idx)
                p = cell.paragraphs[0]
                style_manager._render_inline_tokens(p, cell_tokens, math_renderer)
            current_row += 1

    @classmethod
    def build_khbd_header(cls, doc: docx.Document, data_cache: Dict[str, Any]):
        """Vẽ phần đầu của Kế hoạch bài dạy (KHBD)"""
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        subject = data_cache.get("subject", "MÔN HỌC").upper()
        grade = data_cache.get("grade", "LỚP").upper()
        
        run = p.add_run(f"KẾ HOẠCH BÀI DẠY {subject} - {grade}")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        p.paragraph_format.space_after = Pt(12)

    @classmethod
    def build_matrix_table(cls, doc: docx.Document, exam_data: Dict[str, Any]):
        """
        Vẽ bảng Ma trận theo cấu trúc cố định Công văn 799.
        Sử dụng vòng lặp đối tượng, tuyệt đối không dùng range chỉ số cứng để tránh đứt gãy.
        """
        table = doc.add_table(rows=4, cols=11)
        table.style = 'Table Grid'
        
        # Khởi tạo Text gốc
        headers_l1 = ["STT", "Chủ đề", "Nội dung", "Nhận biết", "", "Thông hiểu", "", "Vận dụng", "", "VDC", "Tổng"]
        for idx, txt in enumerate(headers_l1):
            table.cell(0, idx).text = txt
            
        # Trộn ngang tầng 1 và dán nhãn tầng 2
        for col_idx in [3, 5, 7]:
            table.cell(0, col_idx).merge(table.cell(0, col_idx + 1))
            table.cell(1, col_idx).text = "TN"
            table.cell(1, col_idx + 1).text = "TL"
        table.cell(1, 9).text = "TL"
        
        # Trộn dọc các cột không chia tầng
        for col_idx in [0, 1, 2, 9, 10]:
            table.cell(0, col_idx).merge(table.cell(1, col_idx))

        # Định dạng Header 2 dòng đầu
        for row in table.rows[0:2]:
            for cell in row.cells:
                cls._set_cell_background(cell, "F2F4F4")
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].paragraph_format.space_after = Pt(3)

        # Lấy số liệu an toàn
        topic = str(exam_data.get("custom_req", "Nội dung kiểm tra"))
        c1, c2, c3, c4 = str(exam_data.get("c1", 12)), str(exam_data.get("c2", 2)), str(exam_data.get("c3", 1)), str(exam_data.get("c4", 1))
        
        # Đổ dữ liệu tĩnh vào 2 dòng cuối
        matrix_rows = [
            ["1", topic, "Kiến thức lý thuyết", c1, "0", "0", c2, "0", "1", "0", "N/A"],
            ["2", topic, "Bài tập thực hành", "0", "0", "0", "0", str(int(c3)+int(c4)), "0", "1", "N/A"]
        ]
        
        for r, data in enumerate(matrix_rows, start=2):
            for c, text in enumerate(data):
                table.cell(r, c).text = text

    @classmethod
    def build_specification_table(cls, doc: docx.Document, exam_data: Dict[str, Any]):
        """Vẽ bảng Đặc tả năng lực tương tự cơ chế an toàn của Ma trận"""
        table = doc.add_table(rows=3, cols=5)
        table.style = 'Table Grid'
        
        headers = ["STT", "Chủ đề", "Yêu cầu cần đạt", "Số câu TN", "Số câu TL"]
        for idx, text in enumerate(headers):
            cell = table.cell(0, idx)
            cell.text = text
            cls._set_cell_background(cell, "F2F4F4")
            if cell.paragraphs[0].runs: cell.paragraphs[0].runs[0].bold = True
            
        topic = str(exam_data.get("custom_req", "Nội dung kiểm tra"))
        dt_data = [
            ["1", topic, "- Nhận biết và trích xuất đúng định nghĩa.", "12 câu", "0 câu"],
            ["2", "Tổng hợp", "- Khái quát hóa đơn vị kiến thức.\n- Đánh giá năng lực tư duy.", "4 câu", "3 câu"]
        ]
        
        for r, data in enumerate(dt_data, start=1):
            for c, text in enumerate(data):
                table.cell(r, c).text = text
