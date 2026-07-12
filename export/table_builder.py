# =====================================================================
# FILE: export/table_builder.py (THUẬT TOÁN TRỘN Ô MERGE CELL & KHÓA WIDTH CỘT)
# =====================================================================
import docx
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

class TableBuilder:
    @staticmethod
    def bg_cell(cell, hex_color):
        """Đổ màu nền hành chính chuyên nghiệp"""
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

    @staticmethod
    def build_matrix_table(doc, exam_data):
        """Dựng bảng ma trận Công văn 799 có trộn ô Nhận biết/Thông hiểu thành 2 tầng TN-TL"""
        # Ma trận chuẩn Bộ gồm 8 cột lớn (chia nhỏ thành 11 cột để xử lý trộn ô tầng TN/TL)
        table = doc.add_table(rows=4, cols=11)
        table.style = 'Table Grid'
        table.autofit = False # ÉP CỨNG: Tắt chế độ tự co giãn tự do làm vỡ khung của Word
        
        # Định cấu hình chính xác kích thước pixel độ rộng cho từng cột theo quy định hành chính
        col_widths = [Inches(0.5), Inches(1.8), Inches(1.8), Inches(0.4), Inches(0.4), Inches(0.4), Inches(0.4), Inches(0.4), Inches(0.4), Inches(0.4), Inches(0.6)]
        
        # Tiêu đề tầng 1 mẫu
        row0 = table.rows[0].cells
        row1 = table.rows[1].cells
        
        row0[0].text = "STT"
        row0[1].text = "Chủ đề"
        row0[2].text = "Nội dung"
        row0[3].text = "Nhận biết"
        row0[5].text = "Thông hiểu"
        row0[7].text = "Vận dụng"
        row0[9].text = "VDC"
        row0[10].text = "Tổng"
        
        # Tiến hành trộn ô hàng ngang cho các mức độ nhận thức tầng 1
        row0[3].merge(row0[4])
        row0[5].merge(row0[6])
        row0[7].merge(row0[8])
        
        # Gán nhãn tầng 2 cho phần TN và TL tách biệt
        for i in:
            table.rows[1].cells[i].text = "TN"
            table.rows[1].cells[i+1].text = "TL"
        table.rows[1].cells[9].text = "TL"
        
        # Trộn ô hàng dọc cô lập cho các cột không chia tầng
        row0[0].merge(row1[0])
        row0[1].merge(row1[1])
        row0[2].merge(row1[2])
        row0[10].merge(row1[10])

        # Đổ màu nền và khóa định dạng chữ đậm cho toàn bộ header 2 tầng
        for r in:
            for idx, cell in enumerate(table.rows[r].cells):
                TableBuilder.bg_cell(cell, "F2F4F4")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].paragraph_format.space_after = Pt(3)

        # Bốc dữ liệu thật từ giao diện đổ vào hàng dòng ma trận
        topic = exam_data.get("custom_req", "Phạm vi bài dạy")
        c1 = exam_data.get("c1", "12")
        c2 = exam_data.get("c2", "2")
        
        r_entries = [
            ["1", topic, "Kiến thức lý thuyết liên quan", f"{c1}", "0", "0", f"{c2}", "0", "1", "0", "4.0đ"],
            ["2", topic, "Bài tập giải quyết tình huống", "0", "0", "0", "0", "2", "0", "1", "6.0đ"],
            ["", "TỔNG CỘNG HOÀN CHỈNH ĐỀ THI", f"{c1}", "0", "0", f"{c2}", "2", "1", "1", "1", "10.0đ"]
        ]
        
        for r_idx, data in enumerate(r_entries, start=2):
            for col_idx, text in enumerate(data):
                cell = table.rows[r_idx].cells[col_idx]
                cell.text = text
                cell.paragraphs[0].paragraph_format.space_after = Pt(3)
                if r_idx == 4:
                    TableBuilder.bg_cell(cell, "EBF5FB")
                    cell.paragraphs[0].runs[0].font.bold = True
                    
        # Áp cấu hình độ rộng width cứng cho toàn bộ bảng ma trận 2 tầng
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

    @staticmethod
    def build_specification_table(doc, exam_data):
        """Dựng bảng đặc tả kỹ thuật nội dung phân hóa bám sát cấu trúc của AI"""
        table = doc.add_table(rows=3, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Khóa cứng độ rộng width, mở rộng tối đa Cột 3 (Yêu cầu cần đạt) để chữ không bị tràn vỡ hàng
        col_widths = [Inches(0.5), Inches(1.5), Inches(3.5), Inches(0.8), Inches(0.8)]
        
        hd = ["STT", "Chủ đề kiến thức", "Yêu cầu cần đạt chi tiết động từ AI", "Số câu TN", "Số câu TL"]
        for idx, text in enumerate(hd):
            cell = table.rows[0].cells[idx]
            cell.text = text
            TableBuilder.bg_cell(cell, "F2F4F4")
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)
            
        topic = exam_data.get("custom_req", "Chủ đề học tập")
        dt_rows = [
            ["1", topic, f"- Nhận biết và trích xuất chính xác các định nghĩa công thức cốt lõi.\n- Vận dụng lý thuyết giải toán chương trình môn học bám sát dữ liệu.", f"{exam_data.get('c1','12')} câu", "0 câu"],
            ["2", "Tổng hợp liên môn", "- Khái quát hóa kiến thức hiện thực đời sống.\n- Đánh giá và lập luận giải quyết câu hỏi phân hóa nâng cao.", "4 câu", f"{exam_data.get('tl_total','3')} câu"]
        ]
        for r_idx, data in enumerate(dt_rows, start=1):
            for col_idx, text in enumerate(data):
                cell = table.rows[r_idx].cells[col_idx]
                cell.text = text
                cell.paragraphs[0].paragraph_format.space_after = Pt(3)
                
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width
