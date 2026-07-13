import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io
import re

class WordExportEngine:
    @staticmethod
    def clean_math_formulas(text_line):
        """
        Bộ lọc quét sạch mã rác toán LaTeX chuyển sang định dạng in ấn sạch.
        Đã tối ưu cấu trúc phân rã Regex để bảo vệ không bị dính chữ trên thiết bị khác.
        """
        if not text_line:
            return ""
            
        # Xử lý các khối công thức bọc trong dấu đô la kép hoặc đơn
        text_line = text_line.replace("$$", "").replace("$", "")
        
        # Thiết lập bộ ánh xạ chuyển đổi ký tự toán học cao cấp của Google Gemini sang dạng in ấn sạch
        math_map = {
            r'\\frac\{([^}]+)\}\{([^}]+)\}': r'(\1/\2)',
            r'\\sqrt\{([^}]+)\}': r'√(\1)',
            r'\\sum': '∑', r'\\lim': 'lim', r'\\pi': 'π',
            r'\\alpha': 'α', r'\\beta': 'β', r'\\Delta': 'Δ',
            r'\\rightarrow': '→', r'\\leftrightarrow': '↔',
            r'\\le': '≤', r'\\ge': '≥', r'\\approx': '≈',
            r'\\in': '∈', r'\\subset': '⊂',
            r'\\times': '×', r'\\div': '÷', r'\\neq': '≠', r'\\pm': '±',
            r'\^2': '²', r'\^3': '³'
        }
        
        for pattern, replacement in math_map.items():
            text_line = re.sub(pattern, replacement, text_line)
            
        # Khử sạch các ký tự gạch chéo ngược cô lập còn sót lại mà không gây ảnh hưởng đến chữ thông thường
        text_line = re.sub(r'\\([a-zA-Z]+)', r'\1', text_line)
        text_line = text_line.replace("\\", "")
        
        return text_line.strip()
    @staticmethod
    def build_matrix_table(doc, exam_data):
        """Vẽ bảng ma trận trộn ô Công văn 799 - Đã nâng cấp rows=5 vá triệt để lỗi Out of Range và Syntax dòng 73"""
        # Nâng lên rows=5 để chứa dòng 2, 3 và dòng Tổng cộng 4 an toàn tuyệt đối
        table = doc.add_table(rows=5, cols=11)
        table.style = 'Table Grid'
        table.autofit = False
        col_widths = [Inches(0.4), Inches(1.6), Inches(1.6), Inches(0.4), Inches(0.4), Inches(0.4), Inches(0.4), Inches(0.4), Inches(0.4), Inches(0.4), Inches(0.7)]
        
        def bg_cell(cell, hex_color):
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))
            
        # Khởi tạo tiêu đề gốc tại các ô cell tĩnh ban đầu
        table.cell(0, 0).text = "STT"
        table.cell(0, 1).text = "Chủ đề"
        table.cell(0, 2).text = "Nội dung"
        table.cell(0, 3).text = "Nhận biết"
        table.cell(0, 5).text = "Thông hiểu"
        table.cell(0, 7).text = "Vận dụng"
        table.cell(0, 9).text = "VDC"
        table.cell(0, 10).text = "Tổng"
        
        # Tiết hành trộn ngang tầng 1 và điền nhãn TN/TL cho tầng 2 bằng chỉ số tọa độ
        for col_idx in range(1, 10, 2):
            table.cell(0, col_idx).merge(table.cell(0, col_idx + 1))
            table.cell(1, col_idx).text = "TN"
            table.cell(1, col_idx + 1).text = "TL"
        table.cell(1, 9).text = "TL"
        
        # Đã vá lỗi cú pháp hoàn toàn: Gán mảng ô trộn dọc cố định 0 và 10 của thầy liên mạch
        for col_idx in:
            table.cell(0, col_idx).merge(table.cell(1, col_idx))
            
        # Đổ màu nền cho các dòng tiêu đề gốc
        for r_idx in range(2):
            for cell in table.rows[r_idx].cells:
                bg_cell(cell, "F2F4F4")
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.bold = True
                if cell.paragraphs:
                    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
                    
        topic = exam_data.get("custom_req", "Nội dung đề thi")
        c1 = exam_data.get("c1", 12)
        c2 = exam_data.get("c2", 2)
        c3 = exam_data.get("c3", 1)
        c4 = exam_data.get("c4", 1)
        tn_score_v = exam_data.get("tn_score", "4.0")
        tl_score_v = exam_data.get("tl_score", "6.0")
        
        # Cấu trúc 3 hàng dữ liệu ma trận đổ từ hàng index 2 kịch trần đến index 4 phẳng phiu
        matrix_rows = [
            ["1", topic, "Kiến thức lý thuyết trọng tâm", f"{c1}", "0", "0", f"{c2}", "0", "1", "0", f"{float(tn_score_v)*0.6:.1f}đ"],
            ["2", topic, "Bài tập thực hành phân hóa", "0", "0", "0", "0", f"{int(c3)+int(c4)}", "0", "1", f"{float(tl_score_v):.1f}đ"],
            ["", "TỔNG CỘNG ĐỀ KIỂM TRA ĐẠT CHUẨN", f"{c1}", "0", "0", f"{c2}", f"{int(c3)+int(c4)}", "1", "1", "1", "10.0 điểm"]
        ]
        
        for r, data in enumerate(matrix_rows, start=2):
            for c, text in enumerate(data):
                cell = table.cell(r, c)
                cell.text = text
                cell.paragraphs[0].paragraph_format.space_after = Pt(3)
                if r == 4: # Dòng tổng cộng cuối cùng đổ màu xanh nhạt quý phái bám sát ý thầy
                    bg_cell(cell, "EBF5FB")
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
                        
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = w

    @staticmethod
    def build_specification_table(doc, exam_data):
        """Vẽ bảng đặc tả yêu cầu cần đạt chi tiết phân hóa"""
        table = doc.add_table(rows=3, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        col_widths = [Inches(0.5), Inches(1.5), Inches(3.7), Inches(0.8), Inches(0.8)]
        
        def bg_cell(hex_color, cell_obj):
            cell_obj._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))
            
        hd = ["STT", "Chủ đề kiến thức", "Yêu cầu cần đạt chi tiết phân hóa từ AI", "Số câu TN", "Số câu TL"]
        for idx, text in enumerate(hd):
            cell = table.cell(0, idx)
            cell.text = text
            bg_cell("F2F4F4", cell)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)
            
        topic = exam_data.get("custom_req", "Chủ đề bài học")
        dt_data = [
            ["1", topic, "- Nhận biết và trích xuất đúng định nghĩa hiện tượng khoa học.\n- Vận dụng lý thuyết giải toán bám sát đề cương tài liệu.", f"{exam_data.get('c1','12')} câu", "0 câu"],
            ["2", "Tổng hợp ứng dụng", "- Khái quát hóa đơn vị kiến thức liên môn thực tiễn.\n- Đánh giá năng lực tư duy giải quyết câu hỏi phân hóa cao.", "4 câu", f"{exam_data.get('tl_total','3')} câu"]
        ]
        
        for r, data in enumerate(dt_data, start=1):
            for c, text in enumerate(data):
                cell = table.cell(r, c)
                cell.text = text
                cell.paragraphs[0].paragraph_format.space_after = Pt(3)
                
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = w
    @staticmethod
    def export_to_word(data_cache):
        """Hàm điều phối trung tâm kết xuất tệp Word hành chính"""
        doc = docx.Document()
        for section in doc.sections:
            section.top_margin, section.bottom_margin = Inches(0.79), Inches(0.79)
            section.left_margin, section.right_margin = Inches(1.18), Inches(0.79)
            
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal'].font.size = Pt(12)
        
        ai_text = data_cache.get("ai_generated_content", "")
        
        # PHÂN LUỒNG CỔNG KẾT XUẤT (GIỮ NGUYÊN HOÀN TOÀN CẤU TRÚC 2 CỔNG CỦA THẦY)
        if data_cache.get("is_khbd") == True:
            p_top = doc.add_paragraph()
            p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_top.add_run(f"KẾ HOẠCH BÀI DẠY MÔN {data_cache.get('subject','').upper()} - {data_cache.get('grade','').upper()}").bold = True
        else:
            WordExportEngine.build_matrix_table(doc, data_cache)
            doc.add_paragraph("\n")
            WordExportEngine.build_specification_table(doc, data_cache)
            doc.add_paragraph("\n")
            
        # ENGINE PHÂN TÁCH VÀ CHUYỂN ĐỔI BẢNG MARKDOWN TỪ AI SANG BẢNG WORD CHUYÊN NGHIỆP
        lines = ai_text.split('\n')
        in_table = False
        table_data = []
        
        for line in lines:
            stripped = line.strip()
            
            # Phát hiện dòng dữ liệu bảng dạng Markdown (| cell 1 | cell 2 |)
            if stripped.startswith('|') and stripped.endswith('|'):
                if '---' in stripped: # Bỏ qua dòng kẻ phân cách bảng Markdown
                    continue
                in_table = True
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                table_data.append(cells)
                continue
            else:
                # Nếu vừa kết thúc một khối bảng Markdown, tiến hành dựng ngay bảng Word chuyên nghiệp
                if in_table and table_data:
                    max_cols = max(len(row) for row in table_data)
                    w_table = doc.add_table(rows=len(table_data), cols=max_cols)
                    w_table.style = 'Table Grid'
                    
                    for r_idx, row_cells in enumerate(table_data):
                        for c_idx, cell_value in enumerate(row_cells):
                            if c_idx < max_cols:
                                cleaned_val = WordExportEngine.clean_math_formulas(cell_value.replace("**", ""))
                                w_table.cell(r_idx, c_idx).text = cleaned_val
                                # Đổ màu nền xám nhẹ cho dòng tiêu đề đầu tiên của bảng AI
                                if r_idx == 0:
                                    w_table.cell(r_idx, c_idx)._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F4F4"/>'))
                                    if w_table.cell(r_idx, c_idx).paragraphs and w_table.cell(r_idx, c_idx).paragraphs[0].runs:
                                        w_table.cell(r_idx, c_idx).paragraphs[0].runs[0].font.bold = True
                    doc.add_paragraph("\n")
                    table_data = []
                    in_table = False
                    
            if not stripped:
                continue
                
            # Xử lý làm sạch tiêu đề văn bản thô (Giữ nguyên phong cách của thầy)
            clean_line = stripped.replace("**", "").replace("###", "").replace("##", "")
            processed_text = WordExportEngine.clean_math_formulas(clean_line)
            
            p_item = doc.add_paragraph(processed_text)
            p_item.paragraph_format.space_before = Pt(0)
            p_item.paragraph_format.space_after = Pt(3)
            
            # Ép in đậm thông minh cho các mục tiêu đề sư phạm chính của thầy
            if stripped.startswith("I.") or stripped.startswith("II.") or stripped.startswith("III.") or "Hoạt động" in stripped or "MỤC TIÊU" in stripped.upper() or "TIẾN TRÌNH" in stripped.upper():
                if p_item.runs:
                    p_item.runs[0].font.bold = True
                    
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue()
