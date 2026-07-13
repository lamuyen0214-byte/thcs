"""
Module: export/word_images.py
Nhiệm vụ: Render hình ảnh từ Markdown, chèn Logo, vẽ mã QR và xử lý sơ đồ Mermaid.
Bảo mật: Tự động co giãn (Auto-resize) chống tràn lề trang A4. Xử lý an toàn các URL ảnh hỏng.
Yêu cầu: pip install qrcode
"""

import io
import urllib.request
import urllib.parse
import base64
import logging
from typing import Dict, Any, Optional

import docx
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger("WordImageEngine")

class ImageRenderer:
    # Kích thước chiều ngang tối đa an toàn cho trang A4 (Margins: Trái 3cm, Phải 2cm)
    MAX_A4_WIDTH_INCHES = 6.0 

    @classmethod
    def _fetch_image(cls, source: str) -> Optional[io.BytesIO]:
        """Tải ảnh từ URL mạng hoặc đọc từ đường dẫn cục bộ (Local path) một cách an toàn."""
        try:
            if source.startswith('http://') or source.startswith('https://'):
                # Thêm Header User-Agent để tránh bị các host ảnh chặn (403 Forbidden)
                req = urllib.request.Request(source, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req, timeout=10) as response:
                    image_data = response.read()
                    return io.BytesIO(image_data)
            else:
                # Đọc file ảnh từ ổ cứng cục bộ
                with open(source, 'rb') as f:
                    return io.BytesIO(f.read())
        except Exception as e:
            logger.error(f"Không thể truy xuất hình ảnh từ '{source}': {e}")
            return None

    @classmethod
    def _insert_picture_safe(cls, doc: docx.Document, image_stream: io.BytesIO, width_inches: float = None):
        """Chèn ảnh vào Word, tự động canh giữa và giới hạn kích thước chống tràn lề."""
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        
        try:
            # Nếu người dùng không chỉ định kích thước, lấy tối đa chiều ngang an toàn
            safe_width = min(width_inches, cls.MAX_A4_WIDTH_INCHES) if width_inches else cls.MAX_A4_WIDTH_INCHES
            run.add_picture(image_stream, width=Inches(safe_width))
        except Exception as e:
            logger.error(f"Lỗi tương thích định dạng khi chèn ảnh vào Word: {e}")
            p.add_run("[Lỗi: Không thể hiển thị hình ảnh này]")
        finally:
            p.paragraph_format.space_after = docx.shared.Pt(6)

    @classmethod
    def render_image(cls, doc: docx.Document, node: Dict[str, Any]):
        """Render hình ảnh từ cấu trúc Markdown AST: ![alt](url)"""
        url = node.get("url", "").strip()
        alt_text = node.get("alt", "Image")
        
        if not url:
            return

        logger.info(f"Đang render hình ảnh: {url}")
        image_stream = cls._fetch_image(url)
        
        if image_stream:
            cls._insert_picture_safe(doc, image_stream)
        else:
            # Fallback nếu link ảnh chết
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(f"[Hình ảnh: {alt_text}]")
            run.font.italic = True
            run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)

    @classmethod
    def add_logo(cls, doc: docx.Document, logo_path: str, width_inches: float = 1.0):
        """Chèn Logo trường học vào văn bản (Thường dùng ở Header Kế hoạch bài dạy)"""
        image_stream = cls._fetch_image(logo_path)
        if image_stream:
            cls._insert_picture_safe(doc, image_stream, width_inches)

    @classmethod
    def render_mermaid(cls, doc: docx.Document, mermaid_code: str):
        """
        Dịch mã Mermaid sang hình ảnh SVG/PNG chất lượng cao rồi chèn vào Word.
        Sử dụng API của mermaid.ink (Không yêu cầu cài Node.js trên máy chủ).
        """
        try:
            # Mã hóa Mermaid Code thành Base64 theo chuẩn Mermaid Ink API
            encoded_str = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
            img_url = f"https://mermaid.ink/img/{encoded_str}?type=png"
            
            image_stream = cls._fetch_image(img_url)
            if image_stream:
                cls._insert_picture_safe(doc, image_stream, width_inches=5.0)
            else:
                raise ValueError("Không tải được ảnh từ Mermaid API")
        except Exception as e:
            logger.error(f"Lỗi render sơ đồ Mermaid: {e}")
            p = doc.add_paragraph("[Sơ đồ Mermaid không thể hiển thị]")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    @classmethod
    def add_qr_code(cls, doc: docx.Document, data: str, width_inches: float = 1.5):
        """
        Tạo mã QR tự động từ text/URL (Dùng cho việc cấp link phiếu bài tập online).
        Yêu cầu thư viện: qrcode
        """
        try:
            import qrcode
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=4,
            )
            qr.add_data(data)
            qr.make(fit=True)

            img = qr.make_image(fill_color="black", back_color="white")
            
            # Lưu QR vào bộ nhớ đệm (RAM) rồi chèn thẳng vào Word
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)
            
            cls._insert_picture_safe(doc, img_byte_arr, width_inches)
        except ImportError:
            logger.error("Chưa cài đặt thư viện 'qrcode'. Vui lòng chạy: pip install qrcode")
        except Exception as e:
            logger.error(f"Lỗi khi tạo mã QR: {e}")
