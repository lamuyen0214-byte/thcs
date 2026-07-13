"""
Module: export/word_styles.py
Nhiệm vụ: Quản lý định dạng chuẩn (A4, Times New Roman, Margins).
Kết xuất các Node AST cơ bản: Paragraph, Heading, List Item, Code Block.
"""

import re
from typing import Dict, Any, List
import docx
from docx.shared import Pt, Inches, RGBColor

class StyleManager:
    @staticmethod
    def setup_base_styles(doc: docx.Document):
        """
        Thiết lập cấu hình trang A4 và Font chữ mặc định (Times New Roman 13).
        Lề trang chuẩn hành chính: Trên 2cm, Dưới 2cm, Trái 3cm, Phải 2cm.
        """
        # Cấu hình kích thước trang A4
        for section in doc.sections:
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
            section.top_margin = Inches(0.79)    # ~2.0 cm
            section.bottom_margin = Inches(0.79) # ~2.0 cm
            section.left_margin = Inches(1.18)   # ~3.0 cm
            section.right_margin = Inches(0.79)  # ~2.0 cm

        # Cấu hình Font Normal mặc định
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(13)
        normal_style.font.color.rgb = RGBColor(0, 0, 0)
        
        # Khoảng cách đoạn mặc định
        normal_style.paragraph_format.space_after = Pt(6)

    @staticmethod
    def _render_inline_tokens(paragraph, tokens: List[Dict[str, str]], math_renderer: Any):
        """
        Quét mảng token nội dòng. 
        Nếu là Text -> Bóc tách In đậm/In nghiêng.
        Nếu là Math -> Giao cho MathRenderer (OMML).
        """
        if not tokens:
            return

        for token in tokens:
            t_type = token.get("type")
            content = token.get("content", "")

            if t_type == "text":
                # Bóc tách Markdown In đậm (**) và In nghiêng (*)
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', content)
                for part in parts:
                    if not part:
                        continue
                    if part.startswith('**') and part.endswith('**'):
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        # Tránh bắt nhầm dấu * của phép nhân
                        if len(part) > 2 and part[1] != ' ' and part[-2] != ' ':
                            run = paragraph.add_run(part[1:-1])
                            run.italic = True
                        else:
                            paragraph.add_run(part)
                    else:
                        paragraph.add_run(part)

            elif t_type == "inline_math":
                if math_renderer:
                    # Chuyển giao công thức cho hệ thống OMML Microsoft
                    math_renderer.render_inline_math(paragraph, content)
                else:
                    # Fallback an toàn nếu chưa có module Toán
                    paragraph.add_run(f"{content}")

    @classmethod
    def render_paragraph(cls, doc: docx.Document, node: Dict[str, Any], math_renderer: Any):
        """Vẽ đoạn văn bản tiêu chuẩn"""
        p = doc.add_paragraph()
        cls._render_inline_tokens(p, node.get("tokens", []), math_renderer)

    @classmethod
    def render_heading(cls, doc: docx.Document, node: Dict[str, Any], math_renderer: Any):
        """
        Vẽ tiêu đề (Heading). 
        Sử dụng Paragraph giả lập Heading để chống lỗi màu xanh dương của Word mặc định.
        """
        level = node.get("level", 1)
        p = doc.add_paragraph()
        
        cls._render_inline_tokens(p, node.get("tokens", []), math_renderer)
        
        # Ép định dạng đen, to, in đậm
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            if level == 1:
                run.font.size = Pt(16)
            elif level == 2:
                run.font.size = Pt(15)
            else:
                run.font.size = Pt(14)
                
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True

    @classmethod
    def render_list_item(cls, doc: docx.Document, node: Dict[str, Any], math_renderer: Any):
        """Vẽ danh sách dạng Bullet (Chấm tròn) hoặc Number (Đánh số)"""
        style_name = 'List Number' if node.get("style") == "number" else 'List Bullet'
        p = doc.add_paragraph(style=style_name)
        cls._render_inline_tokens(p, node.get("tokens", []), math_renderer)

    @classmethod
    def render_code_block(cls, doc: docx.Document, node: Dict[str, Any]):
        """Vẽ khối mã lập trình (Code block) với font Courier New"""
        content = node.get("content", "")
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        
        run = p.add_run(content)
        run.font.name = 'Courier New'
        run.font.size = Pt(11)
        # Giả lập nền xám bằng màu chữ xám đậm nếu cần thiết (Word API gốc không hỗ trợ bôi nền Block dễ dàng)
        run.font.color.rgb = RGBColor(50, 50, 50)
