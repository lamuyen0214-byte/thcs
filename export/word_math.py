"""
Module: export/word_math.py
Nhiệm vụ: Biên dịch LaTeX sang OMML (Word Equation) và chuẩn hóa Hóa/Lý.
Yêu cầu: pip install latex2mathml lxml
"""

import re
import logging
from typing import Any
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

logger = logging.getLogger("WordMathEngine")

class ScienceNormalizer:
    """
    Bộ chuẩn hóa chuyên biệt cho Hóa học và Vật lý.
    Xử lý trực tiếp các chuỗi thô do AI sinh ra (không bọc trong thẻ Toán học)
    để đảm bảo hiển thị đúng chuẩn Khoa học Tự nhiên.
    """
    
    # Bảng ánh xạ chỉ số dưới (Subscript) và chỉ số trên (Superscript)
    SUB = str.maketrans("0123456789", "₀₁₂₃₄₅₆₇₈₉")
    SUP = str.maketrans("0123456789+-", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻")

    @classmethod
    def normalize_chemistry(cls, text: str) -> str:
        """
        Chuẩn hóa công thức Hóa học. 
        Ví dụ: H2SO4 -> H₂SO₄ | Fe3+ -> Fe³⁺ | SO4^2- -> SO₄²⁻
        """
        if not text:
            return text

        # 1. Xử lý Ion có ghi chú rõ ràng (VD: SO4^2- -> SO₄²⁻)
        # Bắt nhóm phân tử (SO4) và nhóm ion (^2-)
        text = re.sub(r'([A-Za-z]+\d*)\^(\d*[+\-])', 
                      lambda m: m.group(1).translate(cls.SUB) + m.group(2).translate(cls.SUP), 
                      text)

        # 2. Xử lý Ion viết liền (VD: Fe3+ -> Fe³⁺, NH4+ -> NH₄⁺)
        text = re.sub(r'([A-Z][a-z]?)(\d*[+\-])(?!\w)', 
                      lambda m: m.group(1) + m.group(2).translate(cls.SUP), 
                      text)

        # 3. Xử lý phân tử thông thường (VD: H2SO4 -> H₂SO₄, Ca(OH)2 -> Ca(OH)₂)
        text = re.sub(r'([A-Z][a-z]?|\))(\d+)', 
                      lambda m: m.group(1) + m.group(2).translate(cls.SUB), 
                      text)
        return text

    @classmethod
    def normalize_physics(cls, text: str) -> str:
        """
        Chuẩn hóa đơn vị và ký hiệu Vật lý.
        Ví dụ: m/s^2 -> m/s² | kg.m/s^2 -> kg.m/s²
        """
        if not text:
            return text
            
        # Ánh xạ các ký hiệu Hy Lạp phổ biến của Vật lý
        phys_map = {
            r'\Delta': 'Δ', r'\theta': 'θ', r'\lambda': 'λ',
            r'\mu': 'μ', r'\Omega': 'Ω', r'\epsilon': 'ε',
            r'\sigma': 'σ', r'\rho': 'ρ', r'\varphi': 'φ',
        }
        for tex, uni in phys_map.items():
            text = text.replace(tex, uni)
            
        # Chuẩn hóa số mũ của đơn vị (VD: ^2 -> ², ^-1 -> ⁻¹)
        text = re.sub(r'\^(\d+)', lambda m: m.group(1).translate(cls.SUP), text)
        text = re.sub(r'\^\-(\d+)', lambda m: '⁻' + m.group(1).translate(cls.SUP), text)
        
        return text

class MathRenderer:
    """
    Động cơ render công thức Toán học Microsoft Word (OMML).
    """
    
    # 캐시 (Cache) cho XSLT Transformer để tăng tốc độ render
    _xslt_transform = None

    @classmethod
    def _get_xslt_transformer(cls):
        """Khởi tạo bộ chuyển đổi MathML sang OMML (MML2OMML.XSL)"""
        if cls._xslt_transform is not None:
            return cls._xslt_transform
            
        try:
            from lxml import etree
            import urllib.request
            
            # Tải XSLT chuẩn của Microsoft từ Github (Chỉ tải 1 lần)
            # Trong môi trường Production, thầy có thể tải file MML2OMML.XSL về thư mục export/
            # và đọc từ ổ cứng: tree = etree.parse("export/MML2OMML.XSL")
            xslt_url = "https://raw.githubusercontent.com/plastex/plastex/master/plasTeX/Renderers/Docx/MML2OMML.XSL"
            with urllib.request.urlopen(xslt_url) as response:
                xslt_content = response.read()
                
            xslt_tree = etree.XML(xslt_content)
            cls._xslt_transform = etree.XSLT(xslt_tree)
            return cls._xslt_transform
        except ImportError:
            logger.error("Thiếu thư viện 'lxml'. Vui lòng chạy: pip install lxml")
            return None
        except Exception as e:
            logger.error(f"Không thể khởi tạo XSLT Transformer: {e}")
            return None

    @classmethod
    def latex_to_omml(cls, latex_str: str) -> str:
        """Biên dịch chuỗi LaTeX thành mã XML OMML của Word"""
        try:
            import latex2mathml.converter
            from lxml import etree
            
            # 1. LaTeX -> MathML
            mathml = latex2mathml.converter.convert(latex_str)
            
            # 2. MathML -> OMML
            transform = cls._get_xslt_transformer()
            if not transform:
                raise RuntimeError("XSLT Transformer không khả dụng.")
                
            mml_tree = etree.fromstring(mathml)
            omml_tree = transform(mml_tree)
            
            # Trả về chuỗi XML, bỏ qua XML declaration
            omml_str = str(omml_tree)
            omml_str = re.sub(r'<\?xml.*?\?>', '', omml_str).strip()
            return omml_str
            
        except ImportError:
            logger.error("Thiếu thư viện 'latex2mathml'. Vui lòng chạy: pip install latex2mathml lxml")
            return None
        except Exception as e:
            logger.error(f"Lỗi biên dịch LaTeX ({latex_str}): {e}")
            return None

    @classmethod
    def render_inline_math(cls, paragraph: Any, latex_str: str):
        """Chèn công thức Toán vào giữa dòng văn bản (Inline Math)"""
        # Làm sạch chuỗi trước khi dịch
        latex_str = latex_str.strip()
        if not latex_str:
            return

        omml_xml = cls.latex_to_omml(latex_str)
        
        if omml_xml:
            try:
                # Ép mã OMML trực tiếp vào cấu trúc XML của đoạn văn bản Word
                element = parse_xml(omml_xml)
                paragraph._element.append(element)
            except Exception as e:
                logger.error(f"Lỗi chèn OMML Inline: {e}")
                paragraph.add_run(f"[{latex_str}]") # Fallback an toàn
        else:
            # Nếu hệ thống thiếu thư viện, chuyển hướng sang bộ chuẩn hóa Hóa/Lý
            clean_text = ScienceNormalizer.normalize_chemistry(latex_str)
            clean_text = ScienceNormalizer.normalize_physics(clean_text)
            paragraph.add_run(clean_text)

    @classmethod
    def render_display_math(cls, doc: docx.Document, node: dict):
        """Chèn khối Toán học đa dòng (Display Math / Equations)"""
        latex_str = node.get("content", "").strip()
        if not latex_str:
            return

        p = doc.add_paragraph()
        p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Xử lý các khối matrix, cases, aligned
        omml_xml = cls.latex_to_omml(latex_str)
        
        if omml_xml:
            try:
                element = parse_xml(omml_xml)
                p._element.append(element)
            except Exception as e:
                logger.error(f"Lỗi chèn OMML Block: {e}")
                p.add_run(f"[{latex_str}]")
        else:
            p.add_run(latex_str)
